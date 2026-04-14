#!/usr/bin/env python3
"""
╔══════════════════════════════════════════════════════════════════════════════╗
║           ZAP AI DAST AGENT  v4  —  PRODUCTION READY                       ║
║           Enhanced with 7 Intelligence Pillars                              ║
╠══════════════════════════════════════════════════════════════════════════════╣
║  PILLAR 1 │ Authentication Research      — AI log analysis, 401/403 detect  ║
║  PILLAR 2 │ False Positive Reduction     — Manual FPA vs AI FPA comparison  ║
║  PILLAR 3 │ Findings Prioritization      — OWASP mapping, exploitability    ║
║  PILLAR 4 │ Scan Policy Optimization     — AI-driven policy tuning engine   ║
║  PILLAR 5 │ Trend & Pattern Analysis     — Multi-scan history & deltas      ║
║  PILLAR 6 │ Documentation & Evidence     — Professional Word/Excel reports  ║
║  PILLAR 7 │ Validation                   — AI accuracy tracking             ║
╚══════════════════════════════════════════════════════════════════════════════╝

Requirements:
    pip install python-owasp-zap-v2.4 aiohttp requests python-docx openpyxl
    Docker  (ZAP container)
    Python 3.9+
    Ollama running locally (default: qwen2.5:7b)
"""

# ─────────────────────────────────────────────────────────────────────────────
# IMPORTS
# ─────────────────────────────────────────────────────────────────────────────

from __future__ import annotations

import asyncio
import aiohttp
import csv
import hashlib
import json
import logging
import logging.handlers
import os
import re
import socket
import subprocess
import sys
import threading
import time
from collections import Counter, defaultdict
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field, asdict
from datetime import datetime, timezone
UTC = timezone.utc
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from urllib.parse import urlparse
import xml.etree.ElementTree as ET

import requests

# ── Optional ZAP client ───────────────────────────────────────────────────────
try:
    from zapv2 import ZAPv2
except ImportError:
    ZAPv2 = None

# ── Optional Word / Excel ─────────────────────────────────────────────────────
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False


# ─────────────────────────────────────────────────────────────────────────────
# STRUCTURED CONFIGURATION
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class Config:
    """Central configuration — override via environment variables or config.json."""

    # AI
    ollama_url: str      = os.environ.get("OLLAMA_URL",  "http://127.0.0.1:11434")
    model: str           = os.environ.get("OLLAMA_MODEL", "qwen2.5:7b")

    # Target
    target: str          = os.environ.get("ZAP_TARGET",  "https://pentest-ground.com:4280/")

    # ZAP
    zap_port: int        = int(os.environ.get("ZAP_PORT", "8080"))
    zap_container: str   = os.environ.get("ZAP_CONTAINER", "enterprise-zap")
    zap_image: str       = os.environ.get("ZAP_IMAGE", "ghcr.io/zaproxy/zaproxy:stable")

    # Auth
    zest_login_path: str = os.environ.get("ZEST_LOGIN_PATH", "")
    login_url: str       = os.environ.get("ZAP_LOGIN_URL",  "http://testphp.vulnweb.com/login.php")
    post_url: str        = os.environ.get("ZAP_POST_URL",   "http://testphp.vulnweb.com/userinfo.php")
    default_username: str = os.environ.get("ZAP_USERNAME", "username")
    default_password: str = os.environ.get("ZAP_PASSWORD", "password")

    # Scan limits
    max_scan_time: int       = int(os.environ.get("MAX_SCAN_TIME", "180"))
    spider_max_children: int = int(os.environ.get("SPIDER_MAX_CHILDREN", "3"))
    parallelism: int         = int(os.environ.get("PARALLELISM", "3"))
    poll_interval: float     = 10.0
    poll_timeout: float      = 300.0
    reauth_interval: int     = 540

    # AI settings
    ai_concurrency: int  = int(os.environ.get("AI_CONCURRENCY", "10"))
    ai_timeout: int      = int(os.environ.get("AI_TIMEOUT", "600"))
    ai_retries: int      = 2
    ai_cache_file: str   = "ai_cache.json"
    persist_cache: bool  = True

    # Feature toggles
    enable_active_scan: bool  = True
    enable_ai_analysis: bool  = True
    enable_payload_fuzz: bool = False
    enable_auth: bool         = False
    enable_report: bool       = True
    enable_trend_analysis: bool = True
    enable_policy_optimizer: bool = True

    # Reporting
    report_dir: str      = "reports"
    allowed_risks: frozenset = frozenset({"medium", "high"})
    recurring_threshold: int = 3

    # Risk score weights
    w_ai: float   = 0.40
    w_zap: float  = 0.25
    w_conf: float = 0.15
    w_fuzz: float = 0.20

    # Fuzz rate limiting
    fuzz_delay: float        = 0.2
    fuzz_retries: int        = 3
    fuzz_backoff: float      = 2.0

    # Scope
    excluded_extensions: frozenset = frozenset({
        ".css", ".js", ".png", ".jpg", ".jpeg", ".gif", ".svg",
        ".ico", ".woff", ".woff2", ".ttf", ".eot", ".otf",
        ".pdf", ".zip", ".tar", ".gz", ".mp4", ".mp3", ".webm",
        ".map",
    })
    excluded_path_patterns: Tuple[str, ...] = (
        r"/static/", r"/assets/", r"/media/", r"/fonts/",
        r"/images/", r"/img/", r"/dist/", r"/build/",
    )

    @classmethod
    def from_file(cls, path: str = "config.json") -> "Config":
        """Load config from JSON file if it exists."""
        if os.path.exists(path):
            try:
                with open(path) as f:
                    data = json.load(f)
                obj = cls()
                for k, v in data.items():
                    if hasattr(obj, k):
                        setattr(obj, k, v)
                return obj
            except Exception as e:
                logging.warning("Config file load failed: %s — using defaults", e)
        return cls()

    @property
    def zap_proxy(self) -> str:
        return f"http://127.0.0.1:{self.zap_port}"


CFG = Config.from_file()
os.makedirs(CFG.report_dir, exist_ok=True)


# ─────────────────────────────────────────────────────────────────────────────
# PRODUCTION-GRADE LOGGING
# ─────────────────────────────────────────────────────────────────────────────

def setup_logging() -> logging.Logger:
    """Configure rotating file + console logging with structured format."""
    fmt  = "%(asctime)s | %(levelname)-8s | %(name)-20s | %(message)s"
    dfmt = "%Y-%m-%d %H:%M:%S"

    root = logging.getLogger()
    root.setLevel(logging.DEBUG)
    root.handlers.clear()

    # Console handler — INFO+
    ch = logging.StreamHandler(sys.stdout)
    ch.setLevel(logging.INFO)
    ch.setFormatter(logging.Formatter(fmt, dfmt))
    root.addHandler(ch)

    # Rotating file handler — DEBUG+
    log_file = os.path.join(CFG.report_dir, "dast_agent.log")
    fh = logging.handlers.RotatingFileHandler(
        log_file, maxBytes=10 * 1024 * 1024, backupCount=5, encoding="utf-8"
    )
    fh.setLevel(logging.DEBUG)
    fh.setFormatter(logging.Formatter(fmt, dfmt))
    root.addHandler(fh)

    return logging.getLogger("dast")


log = setup_logging()


# ─────────────────────────────────────────────────────────────────────────────
# RUNTIME STATISTICS (thread-safe)
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class RuntimeStats:
    scan_start: Optional[float]            = None
    scan_end: Optional[float]              = None
    before_ai_count: int                   = 0
    after_ai_count: int                    = 0
    before_ai_high: int                    = 0
    after_ai_high: int                     = 0
    before_ai_medium: int                  = 0
    after_ai_medium: int                   = 0
    duplicates_removed: int                = 0
    dead_paths_pruned: int                 = 0
    urls_before_prune: int                 = 0
    urls_after_prune: int                  = 0
    ai_enrichment_start: Optional[float]   = None
    ai_enrichment_end: Optional[float]     = None
    info_findings_count: int               = 0
    recurring_findings_count: int          = 0
    auth_failures: int                     = 0
    auth_401_403_count: int                = 0
    policy_changes_proposed: int           = 0
    false_positives_removed: int           = 0
    manual_fpa_count: int                  = 0
    ai_fpa_count: int                      = 0
    ai_correct_count: int                  = 0
    ai_incorrect_count: int                = 0
    scan_run_id: str                       = field(
        default_factory=lambda: datetime.now(UTC).strftime("%Y%m%dT%H%M%SZ")
    )

    @property
    def total_duration(self) -> float:
        if self.scan_start and self.scan_end:
            return self.scan_end - self.scan_start
        return 0.0

    @property
    def ai_duration(self) -> float:
        if self.ai_enrichment_start and self.ai_enrichment_end:
            return self.ai_enrichment_end - self.ai_enrichment_start
        return 0.0

    def to_dict(self) -> Dict[str, Any]:
        d = asdict(self)
        d["total_duration_s"] = round(self.total_duration, 2)
        d["ai_duration_s"]    = round(self.ai_duration, 2)
        return d


STATS = RuntimeStats()
_stats_lock = threading.Lock()

# ─────────────────────────────────────────────────────────────────────────────
# KNOWLEDGE BASES
# ─────────────────────────────────────────────────────────────────────────────

OWASP_TOP10_2021: Dict[str, str] = {
    "A01": "Broken Access Control",
    "A02": "Cryptographic Failures",
    "A03": "Injection",
    "A04": "Insecure Design",
    "A05": "Security Misconfiguration",
    "A06": "Vulnerable and Outdated Components",
    "A07": "Identification and Authentication Failures",
    "A08": "Software and Data Integrity Failures",
    "A09": "Security Logging and Monitoring Failures",
    "A10": "Server-Side Request Forgery",
}

OWASP_KEYWORD_MAP: Dict[str, Tuple[str, str]] = {
    "sql":              ("A03", "Injection"),
    "xss":              ("A03", "Injection"),
    "cross-site":       ("A03", "Injection"),
    "injection":        ("A03", "Injection"),
    "csrf":             ("A01", "Broken Access Control"),
    "access control":   ("A01", "Broken Access Control"),
    "idor":             ("A01", "Broken Access Control"),
    "path traversal":   ("A01", "Broken Access Control"),
    "directory":        ("A01", "Broken Access Control"),
    "ssl":              ("A02", "Cryptographic Failures"),
    "tls":              ("A02", "Cryptographic Failures"),
    "certificate":      ("A02", "Cryptographic Failures"),
    "auth":             ("A07", "Identification and Authentication Failures"),
    "session":          ("A07", "Identification and Authentication Failures"),
    "cookie":           ("A07", "Identification and Authentication Failures"),
    "misconfiguration": ("A05", "Security Misconfiguration"),
    "header":           ("A05", "Security Misconfiguration"),
    "version":          ("A06", "Vulnerable and Outdated Components"),
    "component":        ("A06", "Vulnerable and Outdated Components"),
    "log":              ("A09", "Security Logging and Monitoring Failures"),
    "ssrf":             ("A10", "Server-Side Request Forgery"),
    "deserialization":  ("A08", "Software and Data Integrity Failures"),
    "integrity":        ("A08", "Software and Data Integrity Failures"),
}

# Auth failure patterns for AI analysis
AUTH_ERROR_PATTERNS = [
    (r"HTTP [45]\d\d",                  "HTTP error code"),
    (r"401|403|407",                    "Auth / proxy rejection"),
    (r"session.*expir",                 "Session expiry"),
    (r"token.*invalid|invalid.*token",  "Invalid token"),
    (r"captcha|robot|challenge",        "Bot detection"),
    (r"locked|too many attempt",        "Account lockout"),
    (r"password.*incorr|invalid.*cred", "Wrong credentials"),
    (r"timeout|timed out",              "Request timeout"),
]


# ─────────────────────────────────────────────────────────────────────────────
# UTILITY FUNCTIONS
# ─────────────────────────────────────────────────────────────────────────────

def ts() -> str:
    """UTC timestamp string for filenames."""
    return datetime.now(UTC).strftime("%Y%m%dT%H%M%SZ")


def ts_human() -> str:
    return datetime.now(UTC).strftime("%Y-%m-%d %H:%M:%S UTC")


def safe_hash(s: str) -> str:
    return hashlib.sha256(s.encode()).hexdigest()


def extract_json(text: str) -> dict:
    """Robustly extract first JSON object from arbitrary text."""
    if not text:
        return {}
    # Direct parse
    try:
        return json.loads(text)
    except Exception:
        pass
    # Strip markdown fences
    cleaned = re.sub(r"```(?:json)?(.*?)```", r"\1", text, flags=re.DOTALL).strip()
    try:
        return json.loads(cleaned)
    except Exception:
        pass
    # Find first {...}
    m = re.search(r"\{.*\}", cleaned, re.DOTALL)
    if m:
        try:
            return json.loads(m.group())
        except Exception:
            pass
    return {}


def severity_int(sev: str) -> int:
    return {"informational": 0, "info": 0, "low": 1,
            "medium": 2, "high": 3, "critical": 4}.get(
        (sev or "").lower(), -1
    )


def risk_colour_hex(risk: str) -> str:
    r = (risk or "").lower()
    if r in ("critical", "high"): return "FFCDD2"
    if r == "medium":             return "FFE0B2"
    if r in ("low", "info", "informational"): return "C8E6C9"
    return "F5F5F5"


# ─────────────────────────────────────────────────────────────────────────────
# AI CACHE
# ─────────────────────────────────────────────────────────────────────────────

AI_CACHE: Dict[str, Any] = {}
_cache_lock = threading.Lock()


def load_ai_cache() -> None:
    global AI_CACHE
    if not CFG.persist_cache or not os.path.exists(CFG.ai_cache_file):
        return
    try:
        with open(CFG.ai_cache_file, encoding="utf-8") as f:
            with _cache_lock:
                AI_CACHE = json.load(f)
        log.info("AI cache loaded ← %s (%d entries)", CFG.ai_cache_file, len(AI_CACHE))
    except Exception as e:
        log.warning("AI cache load failed: %s", e)


def save_ai_cache() -> None:
    if not CFG.persist_cache:
        return
    try:
        with _cache_lock:
            data = dict(AI_CACHE)
        with open(CFG.ai_cache_file, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2)
        log.info("AI cache saved → %s (%d entries)", CFG.ai_cache_file, len(data))
    except Exception as e:
        log.warning("AI cache save failed: %s", e)


# ─────────────────────────────────────────────────────────────────────────────
# SCAN HISTORY (for trend analysis)
# ─────────────────────────────────────────────────────────────────────────────

SCAN_HISTORY_FILE = os.path.join(CFG.report_dir, "scan_history.json")


def load_scan_history() -> List[Dict[str, Any]]:
    if not os.path.exists(SCAN_HISTORY_FILE):
        return []
    try:
        with open(SCAN_HISTORY_FILE, encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return []


def append_scan_history(stats: RuntimeStats, summary: Dict[str, Any]) -> None:
    history = load_scan_history()
    entry = {**stats.to_dict(), **summary, "timestamp": ts_human()}
    history.append(entry)
    # Keep last 50 scans
    history = history[-50:]
    try:
        with open(SCAN_HISTORY_FILE, "w", encoding="utf-8") as f:
            json.dump(history, f, indent=2)
        log.info("Scan history updated → %s (%d entries)", SCAN_HISTORY_FILE, len(history))
    except Exception as e:
        log.warning("Scan history save failed: %s", e)


# ─────────────────────────────────────────────────────────────────────────────
# PILLAR 1 — AUTHENTICATION RESEARCH
# ─────────────────────────────────────────────────────────────────────────────

AUTH_LOG_FILE = os.path.join(CFG.report_dir, "Auth_log.txt")
_auth_lock     = threading.Lock()
AUTH_SESSION: Optional[requests.Session] = None


def _write_auth_entry(level: str, step: str, detail: str,
                      exc: Exception = None) -> None:
    exc_text = f"\n  Exception : {exc}" if exc else ""
    entry = (
        f"\n{'═' * 70}\n"
        f"[{ts_human()}]  {level}\n"
        f"  Step      : {step}\n"
        f"  Detail    : {detail}{exc_text}\n"
        f"  Target    : {CFG.target}\n"
        f"  Login URL : {CFG.login_url}\n"
        f"{'═' * 70}\n"
    )
    try:
        with open(AUTH_LOG_FILE, "a", encoding="utf-8") as fh:
            fh.write(entry)
    except Exception as e:
        log.error("Cannot write auth log: %s", e)


def log_auth_failure(step: str, detail: str = "", exc: Exception = None) -> None:
    with _stats_lock:
        STATS.auth_failures += 1
    _write_auth_entry("❌ AUTH FAILURE", step, detail, exc)
    log.warning("[Auth] FAILURE | %s — %s", step, detail)


def log_auth_success(step: str, detail: str = "") -> None:
    _write_auth_entry("✅ AUTH SUCCESS", step, detail)
    log.info("[Auth] SUCCESS | %s — %s", step, detail)


# ── 1a. Detect 401 / 403 sequences from ZAP message history ─────────────────

def detect_401_403_sequences(zap) -> Dict[str, Any]:
    """
    Scan ZAP HTTP history for repeated 401/403 patterns.
    Returns analysis dict with counts, URLs, and sequence patterns.
    """
    if zap is None:
        return {}

    log.info("[Auth] Scanning ZAP history for 401/403 sequences…")
    auth_errors: List[Dict] = []

    try:
        messages = zap.core.messages(baseurl=CFG.target, count="500") or []
    except Exception as e:
        log.warning("[Auth] Could not fetch ZAP messages: %s", e)
        return {}

    for msg in messages:
        hdr = msg.get("responseHeader", "")
        if not hdr:
            continue
        status_match = re.search(r"HTTP/\S+ (\d+)", hdr)
        if not status_match:
            continue
        code = int(status_match.group(1))
        if code in (401, 403, 407):
            auth_errors.append({
                "code":    code,
                "url":     msg.get("requestHeader", "").split(" ")[1]
                           if " " in msg.get("requestHeader", "") else "",
                "method":  msg.get("requestHeader", "").split(" ")[0],
                "timestamp": msg.get("timestamp", ""),
            })

    with _stats_lock:
        STATS.auth_401_403_count = len(auth_errors)

    if not auth_errors:
        log.info("[Auth] No 401/403 sequences detected")
        return {"count": 0, "errors": []}

    # Group consecutive errors (within 30 s of each other = "sequence")
    url_counts = Counter(e["url"] for e in auth_errors)
    code_counts = Counter(e["code"] for e in auth_errors)

    sequences_path = os.path.join(CFG.report_dir, f"auth_401_403_{ts()}.json")
    result = {
        "total_auth_errors": len(auth_errors),
        "code_breakdown":    dict(code_counts),
        "top_offending_urls": url_counts.most_common(10),
        "raw_errors":        auth_errors[:50],
    }
    try:
        with open(sequences_path, "w", encoding="utf-8") as f:
            json.dump(result, f, indent=2)
    except Exception:
        pass

    log.info("[Auth] 401/403 analysis: %d errors across %d URLs",
             len(auth_errors), len(url_counts))
    return result


# ── 1b. AI-powered auth log pattern analysis ─────────────────────────────────

async def _analyze_auth_log_async(log_content: str) -> Dict[str, Any]:
    """Send auth log to AI for pattern recognition and recommendations."""
    if not log_content.strip():
        return {}

    prompt = f"""
You are a security automation expert analyzing an authentication log from a DAST scanner.

Analyze the log and return ONLY valid JSON:
{{
  "failure_patterns": [
    {{
      "pattern": "<description of failure pattern>",
      "frequency": <int — how many times>,
      "likely_cause": "<root cause>",
      "severity": "<critical|high|medium|low>",
      "recommendation": "<actionable fix>"
    }}
  ],
  "session_expiry_detected": <true|false>,
  "avg_session_duration_estimate": "<e.g. 9 minutes>",
  "dominant_failure_type": "<single most common failure>",
  "auth_health_score": <0-10 where 10=perfectly stable>,
  "executive_summary": "<2-3 sentence summary for report>"
}}

Auth log (last 3000 chars):
{log_content[-3000:]}
"""
    async with aiohttp.ClientSession() as session:
        raw = await _ask_ai(session, prompt)
    return extract_json(raw)


def analyze_auth_log_with_ai() -> Dict[str, Any]:
    """Run AI analysis on the current auth log file."""
    if not os.path.exists(AUTH_LOG_FILE):
        log.info("[Auth] No auth log found — skipping AI analysis")
        return {}
    try:
        with open(AUTH_LOG_FILE, encoding="utf-8") as f:
            content = f.read()
    except Exception as e:
        log.warning("[Auth] Cannot read auth log: %s", e)
        return {}

    try:
        result = asyncio.run(_analyze_auth_log_async(content))
    except Exception as e:
        log.warning("[Auth] AI log analysis failed: %s", e)
        return {}

    if result:
        out_path = os.path.join(CFG.report_dir, f"auth_ai_analysis_{ts()}.json")
        try:
            with open(out_path, "w", encoding="utf-8") as f:
                json.dump(result, f, indent=2)
            log.info("[Auth] AI auth analysis → %s", out_path)
        except Exception:
            pass

    return result


# ── 1c. Correlate auth failures with scan duration ───────────────────────────

def correlate_auth_with_duration(auth_analysis: Dict[str, Any],
                                 seq_analysis: Dict[str, Any]) -> Dict[str, Any]:
    """
    Correlate authentication failure patterns with scan duration impact.
    Returns a structured report section for inclusion in the main report.
    """
    total_auth_errors = seq_analysis.get("total_auth_errors", 0)
    scan_dur          = STATS.total_duration
    reauth_events     = 0  # set by watchdog

    # Estimate overhead: each 401/403 adds ~2s (retry + backoff)
    estimated_auth_overhead = total_auth_errors * 2.0
    overhead_pct = (estimated_auth_overhead / scan_dur * 100) if scan_dur > 0 else 0

    correlation = {
        "total_scan_duration_s":     round(scan_dur, 1),
        "auth_error_count":          total_auth_errors,
        "estimated_auth_overhead_s": round(estimated_auth_overhead, 1),
        "auth_overhead_pct":         round(overhead_pct, 1),
        "auth_health_score":         auth_analysis.get("auth_health_score", "N/A"),
        "session_expiry_detected":   auth_analysis.get("session_expiry_detected", False),
        "dominant_failure":          auth_analysis.get("dominant_failure_type", "None"),
        "impact_assessment": (
            "High — auth failures significantly degraded scan coverage."
            if overhead_pct > 20 else
            "Medium — auth failures caused moderate scan delays."
            if overhead_pct > 5 else
            "Low — auth failures had minimal impact on scan performance."
        ),
    }

    corr_path = os.path.join(CFG.report_dir, f"auth_correlation_{ts()}.json")
    try:
        with open(corr_path, "w", encoding="utf-8") as f:
            json.dump(correlation, f, indent=2)
    except Exception:
        pass

    log.info("[Auth] Duration correlation: %d auth errors = ~%.1fs overhead (%.1f%%)",
             total_auth_errors, estimated_auth_overhead, overhead_pct)
    return correlation


# ─────────────────────────────────────────────────────────────────────────────
# CORE AUTH HELPERS (unchanged from v3, with enhanced logging)
# ─────────────────────────────────────────────────────────────────────────────

def load_credentials_from_script(script_path: str = None) -> Tuple[str, str]:
    path = script_path or CFG.zest_login_path
    username = password = None

    if not path or not os.path.exists(path):
        msg = f"Zest script not found at '{path}' — using defaults"
        log.warning("[Auth] %s", msg)
        log_auth_failure("load_credentials", msg)
        return CFG.default_username, CFG.default_password

    try:
        with open(path) as f:
            data = json.load(f)
        stmts = data.get("statements", [])

        # Method 1 — browser SendKeys
        send_keys = sorted(
            [s for s in stmts
             if s.get("elementType") == "ZestClientElementSendKeys"
             and s.get("value") and not s["value"].startswith("{{")],
            key=lambda x: x.get("index", 0),
        )
        if len(send_keys) >= 1: username = send_keys[0]["value"]
        if len(send_keys) >= 2: password = send_keys[1]["value"]

        # Method 2 — HTTP form data
        if not username or not password:
            for stmt in stmts:
                for pair in stmt.get("data", "").split("&"):
                    if "=" not in pair: continue
                    k, _, v = pair.partition("=")
                    k = k.strip().lower()
                    v = v.strip()
                    if v.startswith("{{"): continue
                    if k in ("username", "user", "uname", "login") and not username:
                        username = v
                    if k in ("password", "pass", "passwd", "pwd") and not password:
                        password = v

        # Method 3 — script variables
        if not username or not password:
            for var in data.get("parameters", {}).get("variables", []):
                n, v = var.get("name", "").lower(), var.get("value", "")
                if v.startswith("{{"): continue
                if n in ("username", "user", "uname") and not username: username = v
                if n in ("password", "pass", "passwd") and not password: password = v

    except Exception as e:
        log_auth_failure("load_credentials", f"Failed to parse Zest: {path}", e)

    return (username or CFG.default_username), (password or CFG.default_password)


def zap_get(endpoint: str, params: Dict = None) -> Dict:
    try:
        r = requests.get(f"{CFG.zap_proxy}/JSON/{endpoint}",
                         params=params or {}, timeout=10)
        return r.json()
    except Exception as e:
        log.error("ZAP API [%s]: %s", endpoint, e)
        return {}


def auth_cleanup() -> None:
    for ctx in zap_get("context/view/contextList").get("contextList", []):
        if ctx == "vulnweb":
            zap_get("context/action/removeContext", {"contextName": ctx})
    zap_get("script/action/remove", {"scriptName": "login_script"})


def auth_create_context() -> Optional[str]:
    existing = zap_get("context/view/contextList").get("contextList", [])
    if "vulnweb" in existing:
        return zap_get("context/view/context",
                       {"contextName": "vulnweb"}).get("context", {}).get("id")
    res    = zap_get("context/action/newContext", {"contextName": "vulnweb"})
    ctx_id = res.get("contextId")
    if not ctx_id:
        log_auth_failure("auth_create_context", f"No contextId returned: {res}")
    return ctx_id


def auth_include_in_context(context_id: str) -> bool:
    res = zap_get("context/action/includeInContext", {
        "contextName": "vulnweb",
        "regex": f"{CFG.target.rstrip('/')}.*",
    })
    ok = res.get("Result") == "OK"
    if not ok:
        log_auth_failure("auth_include_in_context", str(res))
    return ok


def auth_load_script(script_path: str = None) -> bool:
    path    = script_path or CFG.zest_login_path
    scripts = zap_get("script/view/listScripts").get("scripts", [])
    if any(s.get("name") == "login_script" for s in scripts):
        return True
    res = zap_get("script/action/load", {
        "scriptName": "login_script", "scriptType": "authentication",
        "scriptEngine": "Mozilla Zest", "fileName": path,
    })
    ok = res.get("Result") == "OK"
    if not ok:
        log_auth_failure("auth_load_script", str(res))
    return ok


def auth_set_method(context_id: str, username: str, password: str) -> bool:
    res = zap_get("authentication/action/setAuthenticationMethod", {
        "contextId": context_id,
        "authMethodName": "scriptBasedAuthentication",
        "authMethodConfigParams": (
            f"scriptName=login_script&target={CFG.post_url}"
            f"&username={username}&password={password}"
        ),
    })
    ok = res.get("Result") == "OK"
    if not ok:
        log_auth_failure("auth_set_method", str(res))
    return ok


def auth_set_indicators(context_id: str) -> bool:
    r1 = zap_get("authentication/action/setLoggedInIndicator",
                 {"contextId": context_id, "loggedInIndicatorRegex": "logout"})
    r2 = zap_get("authentication/action/setLoggedOutIndicator",
                 {"contextId": context_id, "loggedOutIndicatorRegex": "(?i)login"})
    ok = r1.get("Result") == "OK" and r2.get("Result") == "OK"
    if not ok:
        log_auth_failure("auth_set_indicators", f"r1={r1}, r2={r2}")
    return ok


def auth_create_user(context_id: str, username: str, password: str) -> Optional[str]:
    res     = zap_get("users/action/newUser",
                      {"contextId": context_id, "name": "testuser"})
    user_id = res.get("userId")
    if not user_id:
        users = zap_get("users/view/usersList",
                        {"contextId": context_id}).get("usersList", [])
        for u in users:
            if u.get("name") == "testuser":
                user_id = u.get("id")
                break
    if not user_id:
        log_auth_failure("auth_create_user", str(res))
        return None
    zap_get("users/action/setAuthenticationCredentials", {
        "contextId": context_id, "userId": user_id,
        "authCredentialsConfigParams": f"username={username}&password={password}",
    })
    zap_get("users/action/setUserEnabled",
            {"contextId": context_id, "userId": user_id, "enabled": True})
    return user_id


def auth_perform_login(username: str, password: str) -> Tuple[bool, Optional[requests.Session]]:
    global AUTH_SESSION
    try:
        session = requests.Session()
        res = session.post(
            CFG.post_url,
            data={"uname": username, "pass": password},
            timeout=10, allow_redirects=True,
            proxies={"http": CFG.zap_proxy, "https": CFG.zap_proxy},
        )
        body = res.text.lower()
        if "logout" in body:
            with _auth_lock:
                AUTH_SESSION = session
            log_auth_success("auth_perform_login", f"HTTP {res.status_code}")
            return True, session

        # Classify failure
        if "invalid" in body or "incorrect" in body or "wrong" in body:
            reason = "Invalid credentials — server rejected"
        elif res.status_code in (401, 403):
            reason = f"HTTP {res.status_code} — access denied"
        elif "captcha" in body or "robot" in body:
            reason = "CAPTCHA / bot-detection challenge"
        else:
            reason = f"'logout' indicator absent (HTTP {res.status_code})"

        log_auth_failure("auth_perform_login", reason)
        return False, None
    except Exception as e:
        log_auth_failure("auth_perform_login", "Exception during login", e)
        return False, None


def auth_verify_zap_session() -> bool:
    time.sleep(2)
    host   = CFG.target.rstrip("/")
    active = [s for s in zap_get("httpSessions/view/sessions",
                                   {"site": host}).get("sessions", [])
              if s.get("active") == "true"]
    if active:
        log_auth_success("auth_verify_zap_session", f"{len(active)} active sessions")
        return True
    for msg in zap_get("core/view/messages",
                        {"baseurl": CFG.post_url, "count": "5"}).get("messages", []):
        hdr  = msg.get("responseHeader", "")
        code = hdr.split(" ")[1] if " " in hdr else "000"
        if code in ("200", "302") and "logout" in msg.get("responseBody", "").lower():
            log_auth_success("auth_verify_zap_session", f"HTTP {code} in history")
            return True
    log_auth_failure("auth_verify_zap_session", "No active session in ZAP after login")
    return False


# ── Session watchdog ──────────────────────────────────────────────────────────

class SessionWatchdog(threading.Thread):
    def __init__(self, username: str, password: str,
                 stop_event: threading.Event,
                 fail_threshold: int = 2):
        super().__init__(daemon=True, name="SessionWatchdog")
        self.username        = username
        self.password        = password
        self.stop_event      = stop_event
        self.reauth_count    = 0
        self._fail_count     = 0
        self._fail_threshold = fail_threshold

    def _is_logged_in(self) -> bool:
        try:
            with _auth_lock:
                sess = AUTH_SESSION
            r = (sess or requests).get(
                CFG.target, timeout=8, allow_redirects=True,
                proxies={"http": CFG.zap_proxy, "https": CFG.zap_proxy},
            )
            return "logout" in (r.text or "").lower()
        except Exception:
            return False

    def run(self) -> None:
        log.info("[Watchdog] Started (probe every %ds)", CFG.reauth_interval)
        while not self.stop_event.wait(timeout=CFG.reauth_interval):
            if not self._is_logged_in():
                self._fail_count += 1
                if self._fail_count < self._fail_threshold:
                    continue
                log.warning("[Watchdog] Session expired — re-authenticating…")
                log_auth_failure("SessionWatchdog",
                                 f"Probe failed {self._fail_count}× — re-auth triggered")
                ok, _ = auth_perform_login(self.username, self.password)
                if ok:
                    self.reauth_count += 1
                    self._fail_count  = 0
                    log.info("[Watchdog] Re-auth #%d succeeded", self.reauth_count)
                else:
                    log.error("[Watchdog] Re-auth failed — scan may be unauthenticated")
            else:
                self._fail_count = 0
        log.info("[Watchdog] Stopped (total re-auths: %d)", self.reauth_count)


def run_full_auth_setup() -> Dict[str, Any]:
    log.info("═" * 60 + "\n         ZAP AUTHENTICATION SETUP\n" + "═" * 60)
    username, password = load_credentials_from_script()
    results: Dict[str, Any] = {}

    try:
        r       = requests.get(f"{CFG.zap_proxy}/JSON/core/view/version/", timeout=5)
        version = r.json().get("version", "?")
        results["ZAP Running"] = True
        log_auth_success("ZAP Running", f"v{version}")
    except Exception as e:
        log_auth_failure("ZAP Running", "ZAP API unreachable", e)
        return {"success": False, "context_id": None, "user_id": None,
                "username": username, "password": password, "results": results}

    auth_cleanup()

    ctx_id = auth_create_context()
    results["Context Created"]   = ctx_id is not None
    if not ctx_id:
        return {"success": False, "context_id": None, "user_id": None,
                "username": username, "password": password, "results": results}

    results["Target in Context"] = auth_include_in_context(ctx_id)

    if not CFG.zest_login_path or not os.path.exists(CFG.zest_login_path):
        log_auth_failure("Zest Script", f"Missing at '{CFG.zest_login_path}'")
        results["Script Loaded"] = False
        return {"success": False, "context_id": ctx_id, "user_id": None,
                "username": username, "password": password, "results": results}

    results["Script Loaded"]  = auth_load_script()
    if not results["Script Loaded"]:
        return {"success": False, "context_id": ctx_id, "user_id": None,
                "username": username, "password": password, "results": results}

    results["Auth Method Set"] = auth_set_method(ctx_id, username, password)
    results["Indicators Set"]  = auth_set_indicators(ctx_id)

    uid = auth_create_user(ctx_id, username, password)
    results["User Created"] = uid is not None
    if not uid:
        return {"success": False, "context_id": ctx_id, "user_id": None,
                "username": username, "password": password, "results": results}

    login_ok, _          = auth_perform_login(username, password)
    results["Login"]     = login_ok
    results["ZAP Session"] = auth_verify_zap_session()
    results["Auth OK"]   = login_ok and results["ZAP Session"]

    overall = all(v for v in results.values() if isinstance(v, bool))
    if overall:
        log_auth_success("Full Auth Setup", "All steps passed")
    else:
        failed = [k for k, v in results.items() if isinstance(v, bool) and not v]
        log_auth_failure("Full Auth Setup", f"Failed steps: {', '.join(failed)}")

    return {"success": overall, "context_id": ctx_id, "user_id": uid,
            "username": username, "password": password, "results": results}


# ─────────────────────────────────────────────────────────────────────────────
# ZAP INFRASTRUCTURE
# ─────────────────────────────────────────────────────────────────────────────

def start_zap() -> None:
    def port_open(p: int) -> bool:
        with socket.socket() as s:
            return s.connect_ex(("127.0.0.1", p)) == 0

    if port_open(CFG.zap_port):
        log.info("ZAP already running on port %d", CFG.zap_port)
        return

    log.info("Starting ZAP Docker container…")
    try:
        subprocess.run([
            "docker", "run", "-d", "--name", CFG.zap_container,
            "-p", f"{CFG.zap_port}:8080", CFG.zap_image,
            "zap.sh", "-daemon", "-host", "0.0.0.0", "-port", "8080",
            "-config", "api.disablekey=true",
        ], check=False)
    except Exception as e:
        log.warning("docker run failed: %s", e)


def wait_for_zap_api(timeout: int = 180) -> None:
    log.info("Waiting for ZAP API (timeout=%ds)…", timeout)
    deadline = time.time() + timeout
    while time.time() < deadline:
        try:
            r = requests.get(f"{CFG.zap_proxy}/JSON/core/view/version/", timeout=3)
            if r.status_code == 200:
                log.info("ZAP API ready")
                return
        except Exception:
            pass
        time.sleep(2)
    raise TimeoutError(f"ZAP API did not respond within {timeout}s")


def zap_client():
    if ZAPv2 is None:
        return None
    zap = ZAPv2(apikey="", proxies={"http": CFG.zap_proxy, "https": CFG.zap_proxy})
    zap._ZAPv2__base = CFG.zap_proxy
    return zap


def zap_force_discover(zap, url: str, wait_seconds: int = 6) -> bool:
    if not url or zap is None:
        return False
    try:
        zap.urlopen(url)
    except Exception:
        pass
    try:
        spid = zap.spider.scan(url)
        t0   = time.time()
        while time.time() - t0 < 4:
            try:
                if int(zap.spider.status(spid)) >= 100:
                    break
            except Exception:
                pass
            time.sleep(0.4)
    except Exception:
        pass
    host = urlparse(url).netloc
    end  = time.time() + wait_seconds
    while time.time() < end:
        try:
            if any(host in s for s in zap.core.sites()):
                return True
        except Exception:
            pass
        time.sleep(0.6)
    return False


# ─────────────────────────────────────────────────────────────────────────────
# SCOPE & FILTERING
# ─────────────────────────────────────────────────────────────────────────────

def is_in_scope(url: str) -> bool:
    if not url:
        return False
    path = urlparse(url.lower()).path
    for ext in CFG.excluded_extensions:
        if path.endswith(ext):
            return False
    for pat in CFG.excluded_path_patterns:
        if re.search(pat, path):
            return False
    return True


def filter_urls(urls: List[str]) -> List[str]:
    before  = len(urls)
    scoped  = [u for u in urls if is_in_scope(u)]
    dropped = before - len(scoped)
    log.info("Static asset filter: kept %d / %d (dropped %d)",
             len(scoped), before, dropped)
    return scoped


# ─────────────────────────────────────────────────────────────────────────────
# RATE-LIMITED HTTP HELPER
# ─────────────────────────────────────────────────────────────────────────────

_fuzz_last: Dict[str, float] = {}
_fuzz_lock = threading.Lock()


def throttled_request(method: str, url: str, **kwargs) -> Optional[requests.Response]:
    host = urlparse(url).netloc
    for attempt in range(CFG.fuzz_retries + 1):
        if CFG.fuzz_delay > 0:
            with _fuzz_lock:
                wait = CFG.fuzz_delay - (time.time() - _fuzz_last.get(host, 0))
                if wait > 0:
                    time.sleep(wait)
                _fuzz_last[host] = time.time()
        try:
            resp = requests.request(method, url, allow_redirects=True,
                                    timeout=15, **kwargs)
        except Exception as e:
            log.warning("HTTP error [%s %s]: %s", method, url, e)
            return None
        if resp.status_code in (429, 503) and attempt < CFG.fuzz_retries:
            backoff = CFG.fuzz_delay * (CFG.fuzz_backoff ** attempt) + 1
            log.warning("Rate-limited (%d) on %s — backoff %.1fs",
                        resp.status_code, host, backoff)
            time.sleep(backoff)
            continue
        return resp
    return None


# ─────────────────────────────────────────────────────────────────────────────
# SPIDER & ACTIVE SCAN
# ─────────────────────────────────────────────────────────────────────────────

def spider_target(zap, target: str,
                  context_name: str = None,
                  user_id: str = None) -> List[str]:
    if zap is None:
        log.warning("ZAP client unavailable — skipping spider")
        return []
    log.info("Spidering: %s", target)
    zap.urlopen(target)
    try:
        spid = (zap.spider.scan_as_user(context_name, user_id, target,
                                        maxchildren=CFG.spider_max_children)
                if context_name and user_id
                else zap.spider.scan(target, maxchildren=CFG.spider_max_children))
    except Exception:
        spid = zap.spider.scan(target, maxchildren=CFG.spider_max_children)

    while True:
        try:
            if int(zap.spider.status(spid)) >= 100:
                break
        except Exception:
            pass
        log.info("Spider: %s%%", zap.spider.status(spid))
        time.sleep(2)

    urls = zap.core.urls()
    log.info("Spider complete — %d URLs discovered", len(urls))
    return urls


def run_parallel_scans(zap, urls: List[str],
                       context_name: str = None,
                       user_id: str = None) -> List[str]:
    if not CFG.enable_active_scan or zap is None:
        log.info("Active scan disabled")
        return []

    urls = list(dict.fromkeys(urls))
    scan_map: Dict[str, Dict] = {}
    start_mono = time.monotonic()
    log.info("Launching active scans (%d URLs, %d workers)", len(urls), CFG.parallelism)

    with ThreadPoolExecutor(max_workers=CFG.parallelism) as ex:
        futures = {ex.submit(lambda u=u: zap.ascan.scan(u, recurse=True), u): u
                   for u in urls}
        for fut in as_completed(futures):
            url = futures[fut]
            try:
                sid = fut.result()
            except Exception as e:
                log.error("Scan start failed for %s: %s", url, e)
                sid = None
            if sid:
                scan_map[str(sid)] = {"url": url, "start": time.monotonic()}
                log.info("Scan %s started → %s", sid, url)
            else:
                log.warning("No scan ID for %s — skipped", url)

    active = dict(scan_map)
    while active:
        elapsed = time.monotonic() - start_mono
        for sid, meta in list(active.items()):
            try:
                pct = int(zap.ascan.status(sid))
            except Exception:
                pct = None
            log.info("Scan %s (%s): %s%%", sid, meta["url"],
                     pct if pct is not None else "?")
            if pct is not None and pct >= 100:
                active.pop(sid, None)
        if elapsed > CFG.max_scan_time:
            log.warning("MAX_SCAN_TIME reached — stopping remaining scans")
            for sid in list(active):
                try: zap.ascan.stop(sid)
                except Exception: pass
            break
        time.sleep(CFG.poll_interval)

    return list(scan_map.keys())


# ─────────────────────────────────────────────────────────────────────────────
# AI CORE
# ─────────────────────────────────────────────────────────────────────────────

async def _ask_ai(session: aiohttp.ClientSession, prompt: str,
                  cache_key: str = None) -> str:
    """Single AI call with caching, retries, and exponential backoff."""
    if cache_key:
        with _cache_lock:
            if cache_key in AI_CACHE:
                return AI_CACHE[cache_key]

    url = f"{CFG.ollama_url}/api/generate"
    for attempt in range(CFG.ai_retries):
        try:
            async with session.post(
                url,
                json={"model": CFG.model, "prompt": prompt,
                      "stream": False, "format": "json",
                      "options": {"temperature": 0}},
                timeout=aiohttp.ClientTimeout(total=CFG.ai_timeout),
            ) as resp:
                d = await resp.json()
                result = d.get("response", "") or d.get("output", "") or json.dumps(d)
                if cache_key and result:
                    with _cache_lock:
                        AI_CACHE[cache_key] = result
                return result
        except Exception as e:
            log.warning("AI attempt %d/%d failed: %s", attempt + 1, CFG.ai_retries, e)
            await asyncio.sleep(2 ** attempt)
    return ""


async def _enrich_alert(session: aiohttp.ClientSession, alert: Dict) -> Dict:
    """Request risk score, severity, summary, and code fix from AI."""
    ck = safe_hash(f"enrich::{alert.get('name')}::{alert.get('risk')}")
    prompt = f"""
You are an application security expert. Return ONLY valid JSON with no markdown.
{{
  "risk_score": <int 0-10>,
  "new_severity": "<Low|Medium|High|Critical>",
  "summary": "<1-2 sentence technical explanation>",
  "remediation_snippet": {{
    "language": "<php|python|java|javascript|generic>",
    "code": "<short code snippet demonstrating the fix>",
    "explanation": "<1 sentence why this fix works>"
  }}
}}

Vulnerability:
  Name      : {alert.get('name')}
  URL       : {alert.get('url')}
  Parameter : {alert.get('param')}
  Risk      : {alert.get('risk')}
  Evidence  : {alert.get('evidence')}
"""
    raw    = await _ask_ai(session, prompt, cache_key=ck)
    parsed = extract_json(raw)
    return {
        "risk_score":          parsed.get("risk_score",          0),
        "new_severity":        parsed.get("new_severity",        "Unknown"),
        "summary":             parsed.get("summary",             "No AI output"),
        "remediation_snippet": parsed.get("remediation_snippet", {}),
    }


async def analyze_alerts_async(alerts: List[Dict]) -> List[Dict]:
    if not CFG.enable_ai_analysis:
        return alerts
    sem = asyncio.Semaphore(CFG.ai_concurrency)

    async with aiohttp.ClientSession() as session:
        async def worker(alert: Dict) -> Dict:
            async with sem:
                enriched = await _enrich_alert(session, alert)
                alert["ai_risk_score"]         = enriched["risk_score"]
                alert["ai_new_severity"]        = enriched["new_severity"]
                alert["ai_summary"]             = enriched["summary"]
                alert["ai_remediation_snippet"] = enriched["remediation_snippet"]
                log.debug("AI enriched: %s", alert.get("name"))
                return alert

        return list(await asyncio.gather(*[worker(a) for a in alerts]))


def analyze_alerts(alerts: List[Dict]) -> List[Dict]:
    return asyncio.run(analyze_alerts_async(alerts))


# ─────────────────────────────────────────────────────────────────────────────
# PILLAR 2 — FALSE POSITIVE REDUCTION
# ─────────────────────────────────────────────────────────────────────────────

async def _ai_group_duplicates_async(alerts: List[Dict]) -> List[Dict]:
    if not alerts:
        return alerts
    names  = list({a.get("name", "") for a in alerts})
    ck     = safe_hash(f"dupes::{sorted(names)}")
    prompt = f"""
You are a security expert. Group the following vulnerability names that are
semantically identical or near-duplicate.

Return ONLY valid JSON:
{{
  "groups": [
    {{
      "canonical_name": "<primary name>",
      "members": ["<name1>", "<name2>"]
    }}
  ]
}}

Vulnerability names:
{json.dumps(names, indent=2)}
"""
    async with aiohttp.ClientSession() as session:
        raw    = await _ask_ai(session, prompt, cache_key=ck)
    parsed = extract_json(raw)
    groups = parsed.get("groups", [])

    name_to_canonical: Dict[str, str] = {}
    for g in groups:
        for member in g.get("members", []):
            name_to_canonical[member] = g.get("canonical_name", member)

    seen: set       = set()
    deduplicated    = []
    removed         = 0

    for alert in alerts:
        name      = alert.get("name", "")
        canonical = name_to_canonical.get(name, name)
        if canonical in seen:
            removed += 1
        else:
            seen.add(canonical)
            alert["canonical_name"]  = canonical
            alert["duplicate_group"] = canonical
            deduplicated.append(alert)

    with _stats_lock:
        STATS.duplicates_removed = removed

    log.info("AI duplicate grouping: %d → %d (removed %d)",
             len(alerts), len(deduplicated), removed)
    return deduplicated


def ai_group_duplicates(alerts: List[Dict]) -> List[Dict]:
    try:
        return asyncio.run(_ai_group_duplicates_async(alerts))
    except Exception as e:
        log.warning("AI duplicate grouping failed: %s", e)
        return alerts


def extract_information_findings(alerts: List[Dict]) -> Tuple[List[Dict], List[Dict]]:
    """Separate informational / disclosure findings from actionable security findings."""
    info_keywords = [
        "information disclosure", "server leaks", "x-powered-by",
        "server version", "directory listing", "debug", "stack trace",
        "internal error", "verbose", "banner", "fingerprint",
        "private ip", "email address", "comments", "source code",
        "web server", "asp.net version", "php version",
    ]
    actionable, informational = [], []
    for alert in alerts:
        name = (alert.get("name") or "").lower()
        risk = (alert.get("risk") or "").lower()
        desc = (alert.get("description") or alert.get("desc") or "").lower()
        is_info = (
            risk in ("informational", "info", "low") and
            any(kw in name or kw in desc for kw in info_keywords)
        )
        if is_info:
            alert["finding_type"] = "informational"
            informational.append(alert)
        else:
            alert["finding_type"] = "security"
            actionable.append(alert)

    with _stats_lock:
        STATS.info_findings_count = len(informational)
    log.info("Informational: %d | Actionable: %d", len(informational), len(actionable))
    return actionable, informational


def identify_recurring_findings(alerts: List[Dict]) -> List[Dict]:
    """Flag findings that appear >= RECURRING_THRESHOLD times."""
    name_counts = Counter(a.get("name", "") for a in alerts)
    recurring_names = {n for n, c in name_counts.items() if c >= CFG.recurring_threshold}

    for alert in alerts:
        name = alert.get("name", "")
        alert["occurrence_count"] = name_counts[name]
        alert["is_recurring"]     = name in recurring_names

    with _stats_lock:
        STATS.recurring_findings_count = len(recurring_names)
    log.info("Recurring findings: %d unique types (≥%d occurrences)",
             len(recurring_names), CFG.recurring_threshold)
    return alerts


# ── Manual FPA vs AI FPA Comparison (new in v4) ───────────────────────────────

def compare_manual_vs_ai_fpa(
    all_alerts: List[Dict],
    manual_fp_names: List[str] = None,
) -> Dict[str, Any]:
    """
    Compare manually identified false positives against AI-identified ones.

    manual_fp_names: list of alert names the analyst marked as FP.
                     If None, look for a manual_fpa.json file in report_dir.

    Returns comparison report and writes it to JSON + CSV.
    """
    # Load manual FPs from file if not provided
    if manual_fp_names is None:
        manual_fp_file = os.path.join(CFG.report_dir, "manual_fpa.json")
        if os.path.exists(manual_fp_file):
            try:
                with open(manual_fp_file) as f:
                    manual_fp_names = json.load(f)
                log.info("[FPA] Loaded %d manual FPs from %s",
                         len(manual_fp_names), manual_fp_file)
            except Exception:
                manual_fp_names = []
        else:
            log.info("[FPA] No manual_fpa.json found — FPA comparison skipped")
            return {}

    manual_set = set(n.lower() for n in (manual_fp_names or []))

    # AI-identified FPs = duplicates + informational findings
    ai_fp_set = set()
    for a in all_alerts:
        name = (a.get("name") or "").lower()
        if a.get("finding_type") == "informational":
            ai_fp_set.add(name)
        if a.get("canonical_name", "").lower() != name:
            ai_fp_set.add(name)

    # Intersection, differences
    both      = manual_set & ai_fp_set
    only_manual = manual_set - ai_fp_set
    only_ai     = ai_fp_set - manual_set

    agreement_pct = (len(both) / max(len(manual_set | ai_fp_set), 1)) * 100

    with _stats_lock:
        STATS.manual_fpa_count = len(manual_set)
        STATS.ai_fpa_count     = len(ai_fp_set)

    report = {
        "manual_fp_count":        len(manual_set),
        "ai_fp_count":            len(ai_fp_set),
        "agreed_fp_count":        len(both),
        "only_manual_fp":         sorted(only_manual),
        "only_ai_fp":             sorted(only_ai),
        "agreement_pct":          round(agreement_pct, 1),
        "missed_by_ai":           sorted(only_manual),   # AI missed these
        "over_flagged_by_ai":     sorted(only_ai),       # AI flagged; human didn't
        "assessment": (
            f"AI and manual analysis agree on {len(both)} false positives "
            f"({agreement_pct:.1f}% agreement). "
            f"AI missed {len(only_manual)} FPs that require manual remediation. "
            f"AI over-flagged {len(only_ai)} findings that may be legitimate issues."
        ),
    }

    out_path = os.path.join(CFG.report_dir, f"fpa_comparison_{ts()}.json")
    try:
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(report, f, indent=2)
        log.info("[FPA] Comparison report → %s", out_path)
    except Exception:
        pass

    return report


# ─────────────────────────────────────────────────────────────────────────────
# PILLAR 3 — FINDINGS PRIORITIZATION
# ─────────────────────────────────────────────────────────────────────────────

async def _ai_rank_findings_async(alerts: List[Dict]) -> List[Dict]:
    if not alerts:
        return alerts

    summaries = [
        {
            "id":       i,
            "name":     a.get("name", ""),
            "risk":     a.get("risk", ""),
            "url":      a.get("url", ""),
            "param":    a.get("param", ""),
            "evidence": (a.get("evidence") or "")[:100],
        }
        for i, a in enumerate(alerts)
    ]
    ck = safe_hash(f"ranking::{[s['name'] for s in summaries]}")
    prompt = f"""
You are a senior penetration tester and risk analyst.
Rank the following security findings by:
  1. exploitation_rank — 1=easiest to exploit (LOWER is more dangerous)
  2. impact_rank       — 1=highest business/data impact (LOWER is worse)

Return ONLY valid JSON:
{{
  "rankings": [
    {{
      "id": <original id>,
      "exploitation_rank": <int>,
      "impact_rank": <int>,
      "exploitation_reason": "<one sentence>",
      "impact_reason": "<one sentence>"
    }}
  ]
}}

Findings:
{json.dumps(summaries, indent=2)}
"""
    async with aiohttp.ClientSession() as session:
        raw    = await _ask_ai(session, prompt, cache_key=ck)
    parsed   = extract_json(raw)
    rankings = {r["id"]: r for r in parsed.get("rankings", [])
                if isinstance(r, dict) and "id" in r}

    for i, alert in enumerate(alerts):
        r = rankings.get(i, {})
        alert["exploitation_rank"]   = r.get("exploitation_rank",   999)
        alert["impact_rank"]         = r.get("impact_rank",         999)
        alert["exploitation_reason"] = r.get("exploitation_reason", "")
        alert["impact_reason"]       = r.get("impact_reason",       "")

    alerts.sort(key=lambda a: (a.get("exploitation_rank", 999) +
                                a.get("impact_rank", 999)))
    log.info("AI ranking complete for %d findings", len(alerts))
    return alerts


def ai_rank_findings(alerts: List[Dict]) -> List[Dict]:
    try:
        return asyncio.run(_ai_rank_findings_async(alerts))
    except Exception as e:
        log.warning("AI ranking failed: %s", e)
        return alerts


def _keyword_owasp_lookup(name: str) -> Tuple[str, str]:
    name_lower = name.lower()
    for kw, (oid, ocat) in OWASP_KEYWORD_MAP.items():
        if kw in name_lower:
            return oid, ocat
    return "A05", "Security Misconfiguration"


async def _ai_map_owasp_async(alerts: List[Dict]) -> List[Dict]:
    if not alerts:
        return alerts
    names = list({a.get("name", "") for a in alerts})
    ck    = safe_hash(f"owasp::{sorted(names)}")
    prompt = f"""
You are a security expert. Map each vulnerability name to its OWASP Top 10 (2021) category.

OWASP Top 10 (2021):
A01-Broken Access Control, A02-Cryptographic Failures, A03-Injection,
A04-Insecure Design, A05-Security Misconfiguration,
A06-Vulnerable and Outdated Components, A07-Identification and Authentication Failures,
A08-Software and Data Integrity Failures, A09-Security Logging and Monitoring Failures,
A10-Server-Side Request Forgery

Return ONLY valid JSON:
{{
  "mappings": [
    {{ "name": "<vuln name>", "owasp_id": "<A01…A10>", "owasp_category": "<name>" }}
  ]
}}

Vulnerabilities:
{json.dumps(names, indent=2)}
"""
    async with aiohttp.ClientSession() as session:
        raw    = await _ask_ai(session, prompt, cache_key=ck)
    parsed   = extract_json(raw)
    mappings = {m["name"]: m for m in parsed.get("mappings", [])
                if isinstance(m, dict) and "name" in m}

    for alert in alerts:
        name = alert.get("name", "")
        if name in mappings:
            alert["owasp_id"]       = mappings[name].get("owasp_id",       "A05")
            alert["owasp_category"] = mappings[name].get("owasp_category", "Security Misconfiguration")
        else:
            oid, ocat = _keyword_owasp_lookup(name)
            alert["owasp_id"]       = oid
            alert["owasp_category"] = ocat

    log.info("OWASP Top 10 mapping complete for %d findings", len(alerts))
    return alerts


def map_owasp_top10(alerts: List[Dict]) -> List[Dict]:
    try:
        return asyncio.run(_ai_map_owasp_async(alerts))
    except Exception as e:
        log.warning("OWASP mapping failed — keyword fallback: %s", e)
        for a in alerts:
            oid, ocat = _keyword_owasp_lookup(a.get("name", ""))
            a["owasp_id"], a["owasp_category"] = oid, ocat
        return alerts


def validate_ai_vs_zap_severity(alerts: List[Dict]) -> Dict[str, Any]:
    """Compare AI severity vs ZAP severity; flag disagreements."""
    agreed = disagreed = escalated = downgraded = 0
    details: List[Dict] = []

    for alert in alerts:
        zap_sev = (alert.get("risk") or "unknown").lower()
        ai_sev  = (alert.get("ai_new_severity") or "unknown").lower()
        zap_val = severity_int(zap_sev)
        ai_val  = severity_int(ai_sev)

        if zap_val == ai_val:
            alert["severity_match"] = True
            alert["severity_delta"] = "agree"
            agreed += 1
        else:
            alert["severity_match"] = False
            delta = ai_val - zap_val
            if delta > 0:
                alert["severity_delta"] = f"AI escalated: {zap_sev} → {ai_sev}"
                escalated += 1
            else:
                alert["severity_delta"] = f"AI downgraded: {zap_sev} → {ai_sev}"
                downgraded += 1
            disagreed += 1
            details.append({
                "name": alert.get("name", ""),
                "url":  alert.get("url",  ""),
                "zap":  zap_sev, "ai": ai_sev,
                "delta": alert["severity_delta"],
            })

    total = len(alerts)
    summary = {
        "total": total, "agreed": agreed, "disagreed": disagreed,
        "escalated": escalated, "downgraded": downgraded,
        "agreement_pct": round(100 * agreed / total, 1) if total else 0,
        "details": details,
    }
    log.info("Severity validation: %.1f%% agreement (%d escalated, %d downgraded)",
             summary["agreement_pct"], escalated, downgraded)

    out = os.path.join(CFG.report_dir, f"severity_validation_{ts()}.json")
    try:
        with open(out, "w", encoding="utf-8") as f:
            json.dump(summary, f, indent=2)
    except Exception:
        pass
    return summary


def compute_priority_score(alert: Dict) -> float:
    ZAP_RISK  = {"informational": 0, "info": 0, "low": 1, "medium": 2, "high": 3}
    ZAP_CONF  = {"false positive": 0, "falsepositive": 0,
                 "low": 1, "medium": 2, "high": 3, "confirmed": 3}
    ai_norm   = min(max(float(alert.get("ai_risk_score", 0)), 0), 10) / 10.0
    zap_risk  = ZAP_RISK.get(str(alert.get("risk", "")).lower(), 0) / 3.0
    zap_conf  = ZAP_CONF.get(str(alert.get("confidence", "")).lower(), 0) / 3.0
    fuzz_hit  = any("reflected" in str(fr.get("evidence", "")).lower() or
                    "server error" in str(fr.get("evidence", "")).lower()
                    for fr in alert.get("fuzz_results", []))
    raw = (CFG.w_ai * ai_norm + CFG.w_zap * zap_risk +
           CFG.w_conf * zap_conf + CFG.w_fuzz * (1.0 if fuzz_hit else 0.0)) * 10
    return round(min(raw, 10.0), 2)


# ─────────────────────────────────────────────────────────────────────────────
# PILLAR 4 — SCAN POLICY OPTIMIZATION
# ─────────────────────────────────────────────────────────────────────────────

class ScanPolicyOptimizer:
    """
    Uses AI insights from scan results to propose concrete scan policy changes.
    Records policy proposals and tracks before/after impact.
    """

    def __init__(self, report_dir: str = CFG.report_dir):
        self.report_dir    = report_dir
        self.proposals: List[Dict] = []

    async def _propose_policy_async(self, alerts: List[Dict],
                                    dead_urls: List[str],
                                    urls: List[str]) -> Dict[str, Any]:
        """Ask AI to propose scan policy changes based on current results."""
        vuln_names  = list({a.get("name", "") for a in alerts})
        owasp_cats  = list({a.get("owasp_category", "") for a in alerts if a.get("owasp_category")})
        url_sample  = urls[:50]

        prompt = f"""
You are a DAST scan optimization expert. Analyze the scan results below and
propose specific scan policy changes to improve efficiency and accuracy.

Return ONLY valid JSON:
{{
  "disable_test_cases": [
    {{
      "test_case": "<ZAP scanner name or category>",
      "reason": "<why irrelevant for this app>",
      "confidence": "<high|medium|low>"
    }}
  ],
  "reduce_crawl_scope": [
    {{
      "pattern": "<URL pattern to exclude>",
      "reason": "<why it wastes scan time>",
      "estimated_time_saving": "<e.g. 5 minutes>"
    }}
  ],
  "increase_focus": [
    {{
      "area": "<test case or URL pattern to focus on>",
      "reason": "<why this area needs more attention>"
    }}
  ],
  "policy_summary": "<2-3 sentence executive summary of recommended changes>",
  "estimated_time_reduction_pct": <int — estimated % reduction in scan time>,
  "estimated_noise_reduction_pct": <int — estimated % reduction in false positives>
}}

Scan context:
  Vulnerabilities found  : {json.dumps(vuln_names[:30])}
  OWASP categories       : {json.dumps(owasp_cats)}
  Dead paths pruned      : {len(dead_urls)}
  URLs in scope          : {len(urls)}
  Sample URLs            : {json.dumps(url_sample)}
  Recurring findings     : {sum(1 for a in alerts if a.get('is_recurring'))}
  Info findings count    : {sum(1 for a in alerts if a.get('finding_type') == 'informational')}
"""
        async with aiohttp.ClientSession() as session:
            raw = await _ask_ai(session, prompt)
        return extract_json(raw)

    def propose_policy_changes(self, alerts: List[Dict],
                               dead_urls: List[str],
                               urls: List[str]) -> Dict[str, Any]:
        if not CFG.enable_policy_optimizer:
            return {}
        try:
            result = asyncio.run(
                self._propose_policy_async(alerts, dead_urls, urls)
            )
        except Exception as e:
            log.warning("[Policy] AI proposal failed: %s", e)
            return {}

        if not result:
            return {}

        with _stats_lock:
            STATS.policy_changes_proposed = (
                len(result.get("disable_test_cases", [])) +
                len(result.get("reduce_crawl_scope", []))
            )

        # Save proposals to file
        out_path = os.path.join(self.report_dir, f"policy_proposals_{ts()}.json")
        try:
            with open(out_path, "w", encoding="utf-8") as f:
                json.dump(result, f, indent=2)
            log.info("[Policy] Proposals saved → %s", out_path)
        except Exception:
            pass

        # Log summary
        disable_count = len(result.get("disable_test_cases", []))
        scope_count   = len(result.get("reduce_crawl_scope", []))
        focus_count   = len(result.get("increase_focus", []))
        log.info("[Policy] Proposals: %d test cases to disable, "
                 "%d scope reductions, %d focus areas",
                 disable_count, scope_count, focus_count)
        log.info("[Policy] Est. time reduction: %d%% | Noise reduction: %d%%",
                 result.get("estimated_time_reduction_pct", 0),
                 result.get("estimated_noise_reduction_pct", 0))

        self.proposals.append(result)
        return result


# ─────────────────────────────────────────────────────────────────────────────
# PILLAR 5 — TREND & PATTERN ANALYSIS
# ─────────────────────────────────────────────────────────────────────────────

class TrendAnalyzer:
    """
    Compares current scan against historical scan runs.
    Detects trends in duration, finding counts, severity distribution.
    """

    def __init__(self):
        self.history = load_scan_history()

    async def _ai_trend_summary_async(self, current: Dict,
                                      history: List[Dict]) -> str:
        """Ask AI to write a trend narrative."""
        if len(history) < 2:
            return "Insufficient scan history for trend analysis."

        prompt = f"""
You are a security metrics analyst. Analyze the scan history and write
a professional trend analysis paragraph (3-5 sentences) covering:
- Direction of finding counts (improving/worsening)
- Changes in scan duration
- Any notable severity trends
- Overall security posture trajectory

Current scan: {json.dumps(current)}
Historical scans (last {len(history)}): {json.dumps(history[-5:], indent=2)}

Return only the narrative text (no JSON, no markdown).
"""
        async with aiohttp.ClientSession() as session:
            return await _ask_ai(session, prompt)

    def analyze(self, stats: RuntimeStats, alerts: List[Dict]) -> Dict[str, Any]:
        """Run full trend analysis and return structured report."""
        if not CFG.enable_trend_analysis:
            return {}

        current_summary = {
            "run_id":         stats.scan_run_id,
            "timestamp":      ts_human(),
            "total_findings": len(alerts),
            "high_findings":  sum(1 for a in alerts
                                  if severity_int(a.get("ai_new_severity")
                                                  or a.get("risk", "")) >= 3),
            "medium_findings": sum(1 for a in alerts
                                   if severity_int(a.get("ai_new_severity")
                                                   or a.get("risk", "")) == 2),
            "scan_duration_s": round(stats.total_duration, 1),
            "duplicates_removed": stats.duplicates_removed,
            "dead_paths_pruned":  stats.dead_paths_pruned,
        }

        if not self.history:
            trend_data = {"status": "First scan — baseline established."}
            trend_narrative = "This is the first recorded scan. A baseline has been established for future trend analysis."
        else:
            # Calculate deltas vs last run
            last = self.history[-1]
            delta_findings  = current_summary["total_findings"] - last.get("total_findings", 0)
            delta_high      = current_summary["high_findings"]  - last.get("high_findings",  0)
            delta_duration  = current_summary["scan_duration_s"] - last.get("scan_duration_s", 0)

            # Duration trend over all scans
            durations = [h.get("scan_duration_s", 0) for h in self.history[-10:]]
            avg_duration = sum(durations) / len(durations) if durations else 0

            # Improvement indicator
            improvements = sum(
                1 for i in range(1, len(self.history))
                if self.history[i].get("total_findings", 0) <
                   self.history[i-1].get("total_findings", 0)
            )

            trend_data = {
                "scans_analyzed":           len(self.history),
                "delta_findings_vs_last":   delta_findings,
                "delta_high_vs_last":       delta_high,
                "delta_duration_vs_last_s": round(delta_duration, 1),
                "avg_scan_duration_10_s":   round(avg_duration, 1),
                "scans_showing_improvement": improvements,
                "trend_direction": (
                    "Improving ↓" if delta_findings < 0 else
                    "Worsening ↑" if delta_findings > 0 else
                    "Stable →"
                ),
            }

            try:
                trend_narrative = asyncio.run(
                    self._ai_trend_summary_async(current_summary, self.history)
                )
            except Exception as e:
                log.warning("[Trend] AI narrative failed: %s", e)
                trend_narrative = (
                    f"Finding count {'decreased' if delta_findings < 0 else 'increased'} "
                    f"by {abs(delta_findings)} compared to the last scan. "
                    f"Scan duration changed by {delta_duration:+.1f}s."
                )

        full_report = {
            "current":         current_summary,
            "trend":           trend_data,
            "narrative":       trend_narrative,
            "history_length":  len(self.history),
        }

        out_path = os.path.join(CFG.report_dir, f"trend_analysis_{ts()}.json")
        try:
            with open(out_path, "w", encoding="utf-8") as f:
                json.dump(full_report, f, indent=2)
            log.info("[Trend] Analysis → %s", out_path)
        except Exception:
            pass

        # Append to history
        append_scan_history(stats, current_summary)
        return full_report


# ─────────────────────────────────────────────────────────────────────────────
# PILLAR 7 — AI VALIDATION TRACKER
# ─────────────────────────────────────────────────────────────────────────────

class ValidationTracker:
    """
    Records where AI was correct vs incorrect for continuous improvement.
    Reads manual review file if it exists; otherwise infers from heuristics.
    """

    MANUAL_REVIEW_FILE = os.path.join(CFG.report_dir, "manual_reviews.json")

    def load_manual_reviews(self) -> List[Dict]:
        """
        Load manual review decisions.
        Expected format: [{"name": "...", "ai_severity": "...",
                           "correct_severity": "...", "is_ai_correct": true/false}]
        """
        if not os.path.exists(self.MANUAL_REVIEW_FILE):
            return []
        try:
            with open(self.MANUAL_REVIEW_FILE) as f:
                return json.load(f)
        except Exception as e:
            log.warning("[Validation] Cannot read manual reviews: %s", e)
            return []

    def evaluate(self, alerts: List[Dict],
                 severity_validation: Dict) -> Dict[str, Any]:
        """
        Produce a validation report:
        - Where manual reviews exist: record AI correct / incorrect
        - Where no manual review: use severity_validation agreement as proxy
        """
        reviews      = self.load_manual_reviews()
        review_index = {r["name"].lower(): r for r in reviews}

        correct   = []
        incorrect = []
        unreviewed = []

        for alert in alerts:
            name   = (alert.get("name") or "").lower()
            ai_sev = (alert.get("ai_new_severity") or "unknown").lower()

            if name in review_index:
                review    = review_index[name]
                is_correct = review.get("is_ai_correct", False)
                entry = {
                    "name":             alert.get("name"),
                    "ai_severity":      ai_sev,
                    "correct_severity": review.get("correct_severity", ""),
                    "verdict":          "correct" if is_correct else "incorrect",
                    "url":              alert.get("url", ""),
                }
                (correct if is_correct else incorrect).append(entry)
            else:
                unreviewed.append({
                    "name": alert.get("name"),
                    "ai_severity": ai_sev,
                    "zap_severity": (alert.get("risk") or "unknown").lower(),
                    "severity_match": alert.get("severity_match", True),
                })

        total_reviewed = len(correct) + len(incorrect)
        accuracy = round(len(correct) / total_reviewed * 100, 1) if total_reviewed else None

        # Use severity agreement as proxy where no manual reviews
        proxy_agreement = severity_validation.get("agreement_pct", 0)

        with _stats_lock:
            STATS.ai_correct_count   = len(correct)
            STATS.ai_incorrect_count = len(incorrect)

        report = {
            "manually_reviewed_count": total_reviewed,
            "ai_correct":              len(correct),
            "ai_incorrect":            len(incorrect),
            "manual_accuracy_pct":     accuracy,
            "proxy_accuracy_pct":      proxy_agreement,
            "unreviewed_count":        len(unreviewed),
            "incorrect_details":       incorrect[:20],
            "correct_details":         correct[:20],
            "recommendations": (
                "AI accuracy is high — consider trusting AI severity without review."
                if (accuracy or proxy_agreement) >= 85 else
                "AI accuracy is moderate — manual review recommended for High/Critical findings."
                if (accuracy or proxy_agreement) >= 70 else
                "AI accuracy needs improvement — validate all High/Critical findings manually."
            ),
            "note": (
                "Accuracy based on manual reviews." if total_reviewed > 0
                else f"No manual reviews found in {self.MANUAL_REVIEW_FILE}. "
                     f"Proxy accuracy ({proxy_agreement:.1f}%) based on ZAP vs AI severity agreement."
            ),
        }

        out_path = os.path.join(CFG.report_dir, f"ai_validation_{ts()}.json")
        try:
            with open(out_path, "w", encoding="utf-8") as f:
                json.dump(report, f, indent=2)
            log.info("[Validation] Report → %s", out_path)
        except Exception:
            pass

        log.info("[Validation] Correct: %d | Incorrect: %d | Unreviewed: %d",
                 len(correct), len(incorrect), len(unreviewed))
        return report


# ─────────────────────────────────────────────────────────────────────────────
# CRAWL SCOPE REDUCTION (Pillar 4 support)
# ─────────────────────────────────────────────────────────────────────────────

async def _ai_dead_paths_async(urls: List[str]) -> List[str]:
    if len(urls) < 5:
        return []
    sample = urls[:100]
    ck     = safe_hash(f"dead::{sorted(sample)}")
    prompt = f"""
You are a security scanning expert. Identify URL patterns that are dead-ends or
irrelevant for security testing (error pages, redirect loops, logout pages,
asset folders, login-only redirects, no dynamic content).

Return ONLY valid JSON:
{{
  "dead_path_patterns": ["<regex pattern1>", ...],
  "dead_urls": ["<exact url1>", ...],
  "reasoning": "<brief explanation>"
}}

Discovered URLs:
{json.dumps(sample, indent=2)}
"""
    async with aiohttp.ClientSession() as session:
        raw  = await _ask_ai(session, prompt, cache_key=ck)
    parsed = extract_json(raw)

    dead_patterns = parsed.get("dead_path_patterns", [])
    dead_urls_ai  = set(parsed.get("dead_urls", []))
    log.info("[Scope] AI dead patterns: %s | Reasoning: %s",
             dead_patterns, parsed.get("reasoning", ""))

    pruned = []
    for url in urls:
        if url in dead_urls_ai or any(
            re.search(p, url, re.IGNORECASE) for p in dead_patterns if p
        ):
            pruned.append(url)
    return pruned


def reduce_crawl_scope(urls: List[str]) -> Tuple[List[str], List[str]]:
    STATS.urls_before_prune = len(urls)
    try:
        dead_urls = asyncio.run(_ai_dead_paths_async(urls))
    except Exception as e:
        log.warning("Dead path detection failed: %s", e)
        dead_urls = []

    dead_set = set(dead_urls)
    active   = [u for u in urls if u not in dead_set]

    STATS.urls_after_prune  = len(active)
    STATS.dead_paths_pruned = len(dead_urls)

    if dead_urls:
        p = os.path.join(CFG.report_dir, f"dead_paths_{ts()}.txt")
        try:
            with open(p, "w") as f:
                f.write(f"Dead paths pruned — {ts_human()}\nTotal: {len(dead_urls)}\n\n")
                for u in dead_urls:
                    f.write(u + "\n")
        except Exception:
            pass

    log.info("Crawl scope: %d → %d (pruned %d dead paths)",
             len(urls), len(active), len(dead_urls))
    return active, dead_urls


# ─────────────────────────────────────────────────────────────────────────────
# MUTATION FUZZING
# ─────────────────────────────────────────────────────────────────────────────

_MUTATIONS = [
    "'", "''", "' OR '1'='1", "' OR 1=1--",
    "<script>alert(1)</script>", "<img src=x onerror=alert(1)>",
    "../", "../../etc/passwd", "%2e%2e%2f",
    "{{7*7}}", "${7*7}", "#{7*7}",
    "; id", "| id", "`id`", "$(id)",
    "\x00", "\r\n", "null", "undefined",
]

TOP_PAYLOADS   = 10
PAYLOAD_TYPES  = ["sql_injection", "xss", "rce", "lfi", "ssti"]


async def _payload_for_alert(session: aiohttp.ClientSession,
                              alert: Dict) -> Dict:
    param = alert.get("param") or ""
    ck    = safe_hash(f"payloads::{alert.get('name')}::{param}")
    prompt = f"""
You are a penetration tester. Generate top {TOP_PAYLOADS} payloads per type.
Return ONLY valid JSON:
{{
  "payloads": {{"<type>": ["payload1", ...]}},
  "hints":    {{"<type>": "inject location: query/body/header/cookie, param name"}}
}}
Target   : {alert.get('name')} @ {alert.get('url')}
Parameter: {param}
Types    : {', '.join(PAYLOAD_TYPES)}
"""
    raw    = await _ask_ai(session, prompt, cache_key=ck)
    parsed = extract_json(raw)
    return {
        "payloads": {t: (parsed.get("payloads", {}).get(t, []) or [])[:TOP_PAYLOADS]
                     for t in PAYLOAD_TYPES},
        "hints":    parsed.get("hints", {}),
    }


def run_payload_fuzz_for_alert(zap, alert: Dict,
                               max_workers: int = 5) -> List[Dict]:
    if not CFG.enable_payload_fuzz:
        return []
    target_url   = alert.get("url")
    param        = alert.get("param") or ""
    payloads_map = alert.get("ai_payloads", {}) or {}
    hints        = alert.get("ai_payload_hints", {}) or {}
    proxies      = {"http": CFG.zap_proxy, "https": CFG.zap_proxy}

    tasks = [(p, ptype, hints.get(ptype, ""))
             for ptype, plist in payloads_map.items() for p in plist]
    original = alert.get("evidence") or param or ""
    for mutant in _MUTATIONS:
        tasks.append((original + mutant, "mutation", f"param:{param}; location:query"))

    if not tasks:
        return []

    def fuzz_task(payload, ptype, hint):
        loc   = "body" if "body" in hint.lower() else \
                "header" if "header" in hint.lower() else "query"
        kw    = {"params": {param or "q": payload}} if loc == "query" else \
                {"data":   {param or "data": payload}}
        resp  = throttled_request("GET" if loc == "query" else "POST",
                                  target_url, proxies=proxies, **kw)
        ev    = "no response"
        if resp:
            ev = ("payload reflected" if payload in (resp.text or "")
                  else f"server error {resp.status_code}" if resp.status_code >= 500
                  else "no evidence")
        return {"type": ptype, "payload": payload, "evidence": ev,
                "status": resp.status_code if resp else None}

    results = []
    with ThreadPoolExecutor(max_workers=max_workers) as ex:
        for fut in as_completed([ex.submit(fuzz_task, *t) for t in tasks]):
            try:
                results.append(fut.result())
            except Exception as e:
                log.warning("Fuzz task error: %s", e)
    return results


# ─────────────────────────────────────────────────────────────────────────────
# PILLAR 6 — DOCUMENTATION: BEFORE/AFTER COMPARISON
# ─────────────────────────────────────────────────────────────────────────────

def build_comparison_rows(stats: RuntimeStats,
                          sev_validation: Dict) -> List[Tuple]:
    """Build the Before vs After comparison table data."""
    return [
        ("Finding Count (total)",
         stats.before_ai_count,  stats.after_ai_count,
         f"{stats.after_ai_count - stats.before_ai_count:+d}"),
        ("High / Critical Findings",
         stats.before_ai_high,   stats.after_ai_high,
         f"{stats.after_ai_high - stats.before_ai_high:+d}"),
        ("Medium Findings",
         stats.before_ai_medium, stats.after_ai_medium,
         f"{stats.after_ai_medium - stats.before_ai_medium:+d}"),
        ("Duplicate Findings Removed",
         "—", stats.duplicates_removed, f"-{stats.duplicates_removed}"),
        ("Informational Findings Separated",
         "—", stats.info_findings_count, f"-{stats.info_findings_count}"),
        ("URLs in Scan Scope",
         stats.urls_before_prune, stats.urls_after_prune,
         f"{stats.urls_after_prune - stats.urls_before_prune:+d}"),
        ("Dead Paths Pruned",
         "—", stats.dead_paths_pruned, f"-{stats.dead_paths_pruned}"),
        ("Manual FPs Identified",
         stats.manual_fpa_count, "—", "—"),
        ("AI FPs Identified",
         "—", stats.ai_fpa_count, "—"),
        ("Auth Error Count (401/403)",
         "—", stats.auth_401_403_count, "—"),
        ("Policy Changes Proposed",
         "—", stats.policy_changes_proposed, "—"),
        ("AI Enrichment Time (s)",
         "—", f"{stats.ai_duration:.1f}", ""),
        ("Total Scan Duration (s)",
         "—", f"{stats.total_duration:.1f}", ""),
        ("Severity Agreement (AI vs ZAP)",
         "—", f"{sev_validation.get('agreement_pct', 0):.1f}%", ""),
        ("AI Correct (validated)",
         "—", stats.ai_correct_count, "—"),
        ("AI Incorrect (validated)",
         "—", stats.ai_incorrect_count, "—"),
    ]


def write_comparison_csv(rows: List[Tuple], path: str) -> None:
    with open(path, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["Metric", "Before AI", "After AI", "Change"])
        w.writerows(rows)
    log.info("Comparison CSV → %s", path)


def write_csv(alerts: List[Dict], path: str) -> None:
    fields = sorted({k for a in alerts for k in a.keys()})
    with open(path, "w", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=fields)
        w.writeheader()
        for a in alerts:
            row = {k: (json.dumps(v) if isinstance(v, (list, dict)) else v)
                   for k, v in a.items()}
            w.writerow(row)
    log.info("CSV → %s", path)


# ─────────────────────────────────────────────────────────────────────────────
# PILLAR 6 — PROFESSIONAL EXCEL REPORT
# ─────────────────────────────────────────────────────────────────────────────

def generate_excel_report(
    alerts: List[Dict],
    info_alerts: List[Dict],
    sev_validation: Dict,
    comparison_rows: List[Tuple],
    fpa_report: Dict,
    policy_report: Dict,
    trend_report: Dict,
    validation_report: Dict,
    path: str = None,
) -> Optional[str]:
    if not XLSX_AVAILABLE:
        log.warning("openpyxl unavailable — Excel report skipped")
        return None

    if path is None:
        path = os.path.join(CFG.report_dir, f"scan_report_{ts()}.xlsx")

    wb = openpyxl.Workbook()

    # ── Style constants ───────────────────────────────────────────────────────
    H_FONT   = Font(name="Calibri", bold=True, color="FFFFFF", size=10)
    H_FILL   = PatternFill("solid", fgColor="1A237E")
    H2_FILL  = PatternFill("solid", fgColor="283593")
    H3_FILL  = PatternFill("solid", fgColor="3949AB")
    ALT_FILL = PatternFill("solid", fgColor="E8EAF6")
    CELL_FONT = Font(name="Calibri", size=9)
    THIN_BORDER = Border(
        left=Side(style="thin",   color="CCCCCC"),
        right=Side(style="thin",  color="CCCCCC"),
        top=Side(style="thin",    color="CCCCCC"),
        bottom=Side(style="thin", color="CCCCCC"),
    )
    THICK_BORDER = Border(
        left=Side(style="medium",   color="1A237E"),
        right=Side(style="medium",  color="1A237E"),
        top=Side(style="medium",    color="1A237E"),
        bottom=Side(style="medium", color="1A237E"),
    )

    def header_cell(ws, row: int, col: int, value: str,
                    fill=None, font_color: str = "FFFFFF") -> None:
        cell = ws.cell(row=row, column=col, value=value)
        cell.font      = Font(name="Calibri", bold=True, color=font_color, size=10)
        cell.fill      = fill or H_FILL
        cell.alignment = Alignment(horizontal="center", vertical="center",
                                   wrap_text=True)
        cell.border    = THICK_BORDER

    def data_cell(ws, row: int, col: int, value,
                  bold: bool = False, fill=None,
                  align: str = "left", wrap: bool = False) -> None:
        cell = ws.cell(row=row, column=col,
                       value=str(value) if value is not None else "")
        cell.font      = Font(name="Calibri", bold=bold, size=9)
        cell.alignment = Alignment(horizontal=align, vertical="top", wrap_text=wrap)
        cell.border    = THIN_BORDER
        if fill:
            cell.fill = fill

    def risk_fill(sev: str) -> Optional[PatternFill]:
        s = (sev or "").lower()
        if s in ("critical", "high"): return PatternFill("solid", fgColor="FFCDD2")
        if s == "medium":             return PatternFill("solid", fgColor="FFE0B2")
        if s in ("low", "info"):      return PatternFill("solid", fgColor="C8E6C9")
        return None

    def section_title(ws, row: int, col: int, title: str) -> None:
        cell = ws.cell(row=row, column=col, value=title)
        cell.font      = Font(name="Calibri", bold=True, size=12, color="1A237E")
        cell.alignment = Alignment(horizontal="left", vertical="center")

    # ══════════════════════════════════════════════════════════════════════════
    # SHEET 1 — ALL FINDINGS
    # ══════════════════════════════════════════════════════════════════════════
    ws1 = wb.active
    ws1.title = "All Findings"
    ws1.sheet_view.showGridLines = False

    section_title(ws1, 1, 1, f"DAST Security Findings  |  Target: {CFG.target}  |  {ts_human()}")
    ws1.merge_cells("A1:R1")

    cols = ["#", "Name", "URL", "Param", "ZAP Risk", "AI Severity",
            "Priority Score", "OWASP ID", "OWASP Category",
            "Exploit Rank", "Impact Rank", "Recurring",
            "Occurrences", "AI Summary", "Fix Language",
            "Fix Explanation", "Sev Match", "Sev Delta"]
    for c, h in enumerate(cols, 1):
        header_cell(ws1, 2, c, h)

    sorted_alerts = sorted(alerts, key=lambda a: a.get("priority_score", 0), reverse=True)
    for i, a in enumerate(sorted_alerts, 1):
        row = i + 2
        rem = a.get("ai_remediation_snippet", {}) or {}
        sev = a.get("ai_new_severity") or a.get("risk", "")
        f   = risk_fill(sev)
        vals = [
            i, a.get("name", ""), a.get("url", ""),
            a.get("param") or "", a.get("risk", ""), sev,
            a.get("priority_score", 0), a.get("owasp_id", ""),
            a.get("owasp_category", ""), a.get("exploitation_rank", ""),
            a.get("impact_rank", ""),
            "✔ Yes" if a.get("is_recurring") else "No",
            a.get("occurrence_count", 1),
            (a.get("ai_summary") or "")[:300],
            rem.get("language", ""), rem.get("explanation", ""),
            "✔" if a.get("severity_match") else "✘",
            a.get("severity_delta", ""),
        ]
        for c, v in enumerate(vals, 1):
            data_cell(ws1, row, c, v,
                      fill=f if c in (5, 6) else
                           ALT_FILL if i % 2 == 0 else None,
                      wrap=(c == 14),
                      align="center" if c in (1, 7, 10, 11, 12, 13, 17) else "left")

    col_widths = [4, 38, 52, 15, 10, 12, 10, 8, 32, 10, 10, 9, 10, 52, 12, 45, 8, 32]
    for c, w in enumerate(col_widths, 1):
        ws1.column_dimensions[get_column_letter(c)].width = w
    ws1.freeze_panes = "A3"

    # ══════════════════════════════════════════════════════════════════════════
    # SHEET 2 — OWASP MAPPING
    # ══════════════════════════════════════════════════════════════════════════
    ws2 = wb.create_sheet("OWASP Mapping")
    ws2.sheet_view.showGridLines = False
    section_title(ws2, 1, 1, "OWASP Top 10 (2021) — Finding Distribution")
    ws2.merge_cells("A1:D1")

    for c, h in enumerate(["OWASP ID", "Category", "Count", "Finding Names"], 1):
        header_cell(ws2, 3, c, h)

    owasp_groups: Dict[str, list] = defaultdict(list)
    for a in alerts:
        owasp_groups[a.get("owasp_id", "A05")].append(a.get("name", "Unknown"))

    for rn, oid in enumerate(sorted(owasp_groups), 4):
        names = owasp_groups[oid]
        ocat  = OWASP_TOP10_2021.get(oid, "Unknown")
        data_cell(ws2, rn, 1, oid, bold=True)
        data_cell(ws2, rn, 2, ocat)
        data_cell(ws2, rn, 3, len(names), align="center")
        data_cell(ws2, rn, 4, "; ".join(set(names))[:250], wrap=True)

    for col_ltr, w in [("A", 10), ("B", 42), ("C", 8), ("D", 65)]:
        ws2.column_dimensions[col_ltr].width = w

    # ══════════════════════════════════════════════════════════════════════════
    # SHEET 3 — BEFORE vs AFTER
    # ══════════════════════════════════════════════════════════════════════════
    ws3 = wb.create_sheet("Before vs After")
    ws3.sheet_view.showGridLines = False
    section_title(ws3, 1, 1, "Before vs After AI Enrichment — Full Comparison")
    ws3.merge_cells("A1:D1")

    for c, h in enumerate(["Metric", "Before AI", "After AI", "Change"], 1):
        header_cell(ws3, 3, c, h)

    for i, row_data in enumerate(comparison_rows):
        rn = i + 4
        for c, v in enumerate(row_data, 1):
            change = str(row_data[3]) if len(row_data) > 3 else ""
            f = None
            if c == 4:
                if change.startswith("-") and change not in ("-0", "-"):
                    f = PatternFill("solid", fgColor="C8E6C9")  # green = reduction = good
                elif change.startswith("+"):
                    f = PatternFill("solid", fgColor="FFCDD2")  # red = increase = bad
            data_cell(ws3, rn, c, v, bold=(c == 1), fill=f,
                      align="center" if c > 1 else "left")

    for col_ltr, w in [("A", 40), ("B", 14), ("C", 14), ("D", 14)]:
        ws3.column_dimensions[col_ltr].width = w

    # ══════════════════════════════════════════════════════════════════════════
    # SHEET 4 — AI RANKINGS
    # ══════════════════════════════════════════════════════════════════════════
    ws4 = wb.create_sheet("AI Rankings")
    ws4.sheet_view.showGridLines = False
    section_title(ws4, 1, 1, "AI-Ranked Findings — Exploitability & Impact")
    ws4.merge_cells("A1:I1")

    rank_cols = ["Exploit Rank", "Impact Rank", "Combined", "Name",
                 "URL", "ZAP Risk", "AI Severity",
                 "Exploitation Reason", "Impact Reason"]
    for c, h in enumerate(rank_cols, 1):
        header_cell(ws4, 3, c, h)

    ranked = sorted(alerts,
                    key=lambda a: (a.get("exploitation_rank", 999) +
                                   a.get("impact_rank", 999)))
    for i, a in enumerate(ranked):
        rn   = i + 4
        comb = (a.get("exploitation_rank", 999) + a.get("impact_rank", 999))
        f    = risk_fill(a.get("ai_new_severity") or a.get("risk"))
        vals = [a.get("exploitation_rank", ""), a.get("impact_rank", ""), comb,
                a.get("name", ""), a.get("url", ""),
                a.get("risk", ""), a.get("ai_new_severity", ""),
                a.get("exploitation_reason", ""), a.get("impact_reason", "")]
        for c, v in enumerate(vals, 1):
            data_cell(ws4, rn, c, v,
                      fill=f if c in (6, 7) else
                           ALT_FILL if i % 2 == 0 else None,
                      wrap=(c in (8, 9)),
                      align="center" if c <= 3 else "left")

    for c, w in enumerate([10, 10, 10, 38, 45, 10, 12, 55, 55], 1):
        ws4.column_dimensions[get_column_letter(c)].width = w
    ws4.freeze_panes = "A4"

    # ══════════════════════════════════════════════════════════════════════════
    # SHEET 5 — INFORMATION FINDINGS
    # ══════════════════════════════════════════════════════════════════════════
    ws5 = wb.create_sheet("Information Findings")
    ws5.sheet_view.showGridLines = False
    section_title(ws5, 1, 1, "Informational Findings — Disclosure & Low Risk")
    ws5.merge_cells("A1:F1")

    for c, h in enumerate(["#", "Name", "URL", "Param", "Risk", "Description"], 1):
        header_cell(ws5, 3, c, h)

    for i, a in enumerate(info_alerts, 1):
        rn = i + 3
        vals = [i, a.get("name", ""), a.get("url", ""),
                a.get("param") or "", a.get("risk", ""),
                (a.get("description") or a.get("desc") or "")[:200]]
        for c, v in enumerate(vals, 1):
            data_cell(ws5, rn, c, v, wrap=(c == 6),
                      fill=ALT_FILL if i % 2 == 0 else None)

    for col_ltr, w in [("A",4),("B",42),("C",52),("D",14),("E",10),("F",65)]:
        ws5.column_dimensions[col_ltr].width = w

    # ══════════════════════════════════════════════════════════════════════════
    # SHEET 6 — FPA COMPARISON
    # ══════════════════════════════════════════════════════════════════════════
    ws6 = wb.create_sheet("FPA Comparison")
    ws6.sheet_view.showGridLines = False
    section_title(ws6, 1, 1, "False Positive Analysis — Manual vs AI Comparison")
    ws6.merge_cells("A1:B1")

    if fpa_report:
        pairs = [
            ("Manual FP Count",       fpa_report.get("manual_fp_count", 0)),
            ("AI FP Count",            fpa_report.get("ai_fp_count", 0)),
            ("Agreed FPs",             fpa_report.get("agreed_fp_count", 0)),
            ("Agreement %",            f"{fpa_report.get('agreement_pct', 0):.1f}%"),
            ("Missed by AI",           len(fpa_report.get("missed_by_ai", []))),
            ("Over-flagged by AI",     len(fpa_report.get("over_flagged_by_ai", []))),
        ]
        header_cell(ws6, 3, 1, "Metric")
        header_cell(ws6, 3, 2, "Value")
        for i, (k, v) in enumerate(pairs, 4):
            data_cell(ws6, i, 1, k, bold=True)
            data_cell(ws6, i, 2, v, align="center")

        rn = len(pairs) + 5
        data_cell(ws6, rn, 1, "Assessment", bold=True)
        data_cell(ws6, rn, 2, fpa_report.get("assessment", ""), wrap=True)

        rn += 2
        for lbl, items, fill_hex in [
            ("Missed by AI (Manual only)",    fpa_report.get("missed_by_ai",       []), "FFCDD2"),
            ("Over-flagged by AI (AI only)",  fpa_report.get("over_flagged_by_ai", []), "FFE0B2"),
        ]:
            data_cell(ws6, rn, 1, lbl, bold=True)
            rn += 1
            for item in items:
                data_cell(ws6, rn, 1, item,
                          fill=PatternFill("solid", fgColor=fill_hex))
                rn += 1
            rn += 1
    else:
        data_cell(ws6, 3, 1, "No manual_fpa.json provided — see README for instructions.")

    ws6.column_dimensions["A"].width = 45
    ws6.column_dimensions["B"].width = 55

    # ══════════════════════════════════════════════════════════════════════════
    # SHEET 7 — SCAN POLICY PROPOSALS
    # ══════════════════════════════════════════════════════════════════════════
    ws7 = wb.create_sheet("Policy Proposals")
    ws7.sheet_view.showGridLines = False
    section_title(ws7, 1, 1, "AI Scan Policy Optimization Proposals")
    ws7.merge_cells("A1:D1")

    if policy_report:
        rn = 3
        data_cell(ws7, rn, 1, "Summary", bold=True)
        data_cell(ws7, rn, 2, policy_report.get("policy_summary", ""), wrap=True)
        ws7.merge_cells(f"B{rn}:D{rn}")

        rn += 1
        for col_ltr, label, value in [
            ("A", "Est. Time Reduction",  f"{policy_report.get('estimated_time_reduction_pct', 0)}%"),
            ("B", "Est. Noise Reduction", f"{policy_report.get('estimated_noise_reduction_pct', 0)}%"),
        ]:
            data_cell(ws7, rn, ord(col_ltr)-64, f"{label}: {value}", bold=True)

        rn += 2
        # Test cases to disable
        disable_items = policy_report.get("disable_test_cases", [])
        header_cell(ws7, rn, 1, "Test Cases to Disable")
        header_cell(ws7, rn, 2, "Reason")
        header_cell(ws7, rn, 3, "Confidence")
        ws7.merge_cells(f"A{rn}:A{rn}")
        rn += 1
        for item in disable_items:
            data_cell(ws7, rn, 1, item.get("test_case", ""))
            data_cell(ws7, rn, 2, item.get("reason", ""), wrap=True)
            data_cell(ws7, rn, 3, item.get("confidence", ""), align="center")
            rn += 1

        rn += 1
        # Scope reductions
        scope_items = policy_report.get("reduce_crawl_scope", [])
        header_cell(ws7, rn, 1, "Crawl Scope Reductions")
        header_cell(ws7, rn, 2, "Reason")
        header_cell(ws7, rn, 3, "Est. Time Saving")
        rn += 1
        for item in scope_items:
            data_cell(ws7, rn, 1, item.get("pattern", ""))
            data_cell(ws7, rn, 2, item.get("reason", ""), wrap=True)
            data_cell(ws7, rn, 3, item.get("estimated_time_saving", ""), align="center")
            rn += 1
    else:
        data_cell(ws7, 3, 1, "Policy optimization not run or no proposals generated.")

    for col_ltr, w in [("A", 42), ("B", 55), ("C", 18), ("D", 18)]:
        ws7.column_dimensions[col_ltr].width = w

    # ══════════════════════════════════════════════════════════════════════════
    # SHEET 8 — TREND ANALYSIS
    # ══════════════════════════════════════════════════════════════════════════
    ws8 = wb.create_sheet("Trend Analysis")
    ws8.sheet_view.showGridLines = False
    section_title(ws8, 1, 1, "Scan Trend & Pattern Analysis — Historical Comparison")
    ws8.merge_cells("A1:B1")

    if trend_report:
        rn = 3
        data_cell(ws8, rn, 1, "AI Trend Narrative", bold=True)
        data_cell(ws8, rn, 2, trend_report.get("narrative", ""), wrap=True)
        ws8.row_dimensions[rn].height = 80

        rn += 2
        trend = trend_report.get("trend", {})
        header_cell(ws8, rn, 1, "Trend Metric")
        header_cell(ws8, rn, 2, "Value")
        rn += 1
        for k, v in trend.items():
            data_cell(ws8, rn, 1, k.replace("_", " ").title())
            data_cell(ws8, rn, 2, v, align="center")
            rn += 1

        rn += 1
        # Historical table
        history = load_scan_history()
        if len(history) > 1:
            header_cell(ws8, rn, 1, "Scan Run ID")
            header_cell(ws8, rn, 2, "Timestamp")
            header_cell(ws8, rn, 3, "Total Findings")
            header_cell(ws8, rn, 4, "High Findings")
            header_cell(ws8, rn, 5, "Duration (s)")
            rn += 1
            for h in history[-10:]:
                data_cell(ws8, rn, 1, h.get("scan_run_id", ""))
                data_cell(ws8, rn, 2, h.get("timestamp", ""))
                data_cell(ws8, rn, 3, h.get("total_findings", ""), align="center")
                data_cell(ws8, rn, 4, h.get("high_findings", ""), align="center")
                data_cell(ws8, rn, 5, h.get("scan_duration_s", ""), align="center")
                rn += 1

    for col_ltr, w in [("A", 35), ("B", 65), ("C", 15), ("D", 15), ("E", 14)]:
        ws8.column_dimensions[col_ltr].width = w

    # ══════════════════════════════════════════════════════════════════════════
    # SHEET 9 — AI VALIDATION
    # ══════════════════════════════════════════════════════════════════════════
    ws9 = wb.create_sheet("AI Validation")
    ws9.sheet_view.showGridLines = False
    section_title(ws9, 1, 1, "AI Accuracy Validation — Correct vs Incorrect Decisions")
    ws9.merge_cells("A1:B1")

    if validation_report:
        rn = 3
        pairs = [
            ("Manually Reviewed",     validation_report.get("manually_reviewed_count", 0)),
            ("AI Correct",            validation_report.get("ai_correct", 0)),
            ("AI Incorrect",          validation_report.get("ai_incorrect", 0)),
            ("Manual Accuracy %",     validation_report.get("manual_accuracy_pct") or "N/A"),
            ("Proxy Accuracy %",      f"{validation_report.get('proxy_accuracy_pct', 0):.1f}%"),
            ("Unreviewed Findings",   validation_report.get("unreviewed_count", 0)),
        ]
        header_cell(ws9, rn, 1, "Metric")
        header_cell(ws9, rn, 2, "Value")
        rn += 1
        for k, v in pairs:
            data_cell(ws9, rn, 1, k, bold=True)
            data_cell(ws9, rn, 2, v, align="center")
            rn += 1

        rn += 1
        data_cell(ws9, rn, 1, "Recommendations", bold=True)
        data_cell(ws9, rn, 2, validation_report.get("recommendations", ""), wrap=True)

        rn += 2
        data_cell(ws9, rn, 1, "Note", bold=True)
        data_cell(ws9, rn, 2, validation_report.get("note", ""), wrap=True)

        incorrect = validation_report.get("incorrect_details", [])
        if incorrect:
            rn += 2
            header_cell(ws9, rn, 1, "Finding Name")
            header_cell(ws9, rn, 2, "AI Severity")
            header_cell(ws9, rn, 3, "Correct Severity")
            rn += 1
            for item in incorrect:
                data_cell(ws9, rn, 1, item.get("name", ""))
                data_cell(ws9, rn, 2, item.get("ai_severity", ""), align="center")
                data_cell(ws9, rn, 3, item.get("correct_severity", ""), align="center")
                rn += 1

    ws9.column_dimensions["A"].width = 40
    ws9.column_dimensions["B"].width = 55
    ws9.column_dimensions["C"].width = 20

    wb.save(path)
    log.info("Excel report → %s", path)
    return path


# ─────────────────────────────────────────────────────────────────────────────
# PILLAR 6 — PROFESSIONAL WORD REPORT
# ─────────────────────────────────────────────────────────────────────────────

async def _ai_executive_summary_async(alerts: List[Dict],
                                       stats: RuntimeStats,
                                       sev_validation: Dict,
                                       trend_report: Dict) -> str:
    high  = sum(1 for a in alerts if severity_int(
        a.get("ai_new_severity") or a.get("risk", "")) >= 3)
    med   = sum(1 for a in alerts if severity_int(
        a.get("ai_new_severity") or a.get("risk", "")) == 2)
    trend = trend_report.get("trend", {}).get("trend_direction", "Unknown")
    owasp = list({a.get("owasp_category") for a in alerts if a.get("owasp_category")})

    prompt = f"""
You are a senior security consultant writing an executive summary for a C-suite DAST scan report.
Write a professional 3-4 paragraph summary covering: overall risk posture, most critical findings,
business impact, remediation priority, and security trend direction.

Target              : {CFG.target}
Total findings      : {len(alerts)}
High / Critical     : {high}
Medium              : {med}
OWASP categories    : {json.dumps(owasp[:8])}
Severity agreement  : {sev_validation.get('agreement_pct', 0):.1f}%
Duplicates removed  : {stats.duplicates_removed}
Trend direction     : {trend}
Auth errors (401/403): {stats.auth_401_403_count}

Return only the paragraph text (no JSON, no markdown headers).
"""
    async with aiohttp.ClientSession() as session:
        return await _ask_ai(session, prompt)

def generate_word_report(
    alerts: List[Dict],
    info_alerts: List[Dict],
    sev_validation: Dict,
    comparison_rows: List[Tuple],
    stats: RuntimeStats,
    auth_analysis: Dict,
    auth_correlation: Dict,
    fpa_report: Dict,
    policy_report: Dict,
    trend_report: Dict,
    validation_report: Dict,
    path: str = None,
) -> Optional[str]:
    if not DOCX_AVAILABLE:
        log.warning("python-docx unavailable — Word report skipped")
        return None

    if path is None:
        path = os.path.join(CFG.report_dir, f"scan_report_{ts()}.docx")

    # Generate AI executive summary
    try:
        exec_summary = asyncio.run(
            _ai_executive_summary_async(alerts, stats, sev_validation, trend_report)
        )
        if not exec_summary or len(exec_summary.strip()) < 30:
            raise ValueError("Empty AI summary")
    except Exception as e:
        log.warning("AI executive summary failed: %s", e)
        exec_summary = (
            f"A DAST scan was performed against {CFG.target}. "
            f"A total of {len(alerts)} unique security findings were identified. "
            "Please review the tables below for details."
        )

    doc = DocxDocument()

    # ── Page setup (US Letter: 8.5 × 11 in) ──────────────────────────────────
    # python-docx requires EMU (English Metric Units).
    # Old raw twip values (12240, 15840, 1080) produced a ~0.013-inch-wide
    # page, corrupting the document layout. Inches() converts correctly.
    section = doc.sections[0]
    section.page_width    = Inches(8.5)   # 8.5 in  → 7,772,400 EMU
    section.page_height   = Inches(11)    # 11 in   → 10,058,400 EMU
    section.left_margin   = Inches(0.75)  # 0.75 in → replaces 1080 twips
    section.right_margin  = Inches(0.75)
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # ... rest of the function remains unchanged
    # Helper functions
    def h(text: str, level: int = 1) -> None:
        p = doc.add_heading(text, level=level)
        colour = RGBColor(0x1A, 0x23, 0x7E) if level == 1 else RGBColor(0x28, 0x35, 0x93)
        for run in p.runs:
            run.font.color.rgb = colour

    def para(text: str, bold: bool = False, size: int = 10,
             style: str = "Normal") -> None:
        p = doc.add_paragraph(style=style)
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.bold = bold

    def shade(cell, hex_color: str) -> None:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def cell_text(cell, text, bold: bool = False, size: int = 9,
                  align: str = "left") -> None:
        cell.text = ""
        p   = cell.paragraphs[0]
        run = p.add_run(str(text) if text is not None else "")
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.bold = bold
        p.alignment = {
            "left":   WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right":  WD_ALIGN_PARAGRAPH.RIGHT,
        }.get(align, WD_ALIGN_PARAGRAPH.LEFT)

    def add_table_with_header(headers: list, header_fill: str = "1A237E") -> object:
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Table Grid"
        hdr_cells = tbl.rows[0].cells
        for i, hdr in enumerate(headers):
            cell_text(hdr_cells[i], hdr, bold=True, align="center")
            shade(hdr_cells[i], header_fill)
            for run in hdr_cells[i].paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        return tbl

    # ── Title page ────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_run  = title_para.add_run("DAST Security Assessment Report")
    title_run.font.size  = Pt(28)
    title_run.font.bold  = True
    title_run.font.name  = "Calibri"
    title_run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        f"Target: {CFG.target}",
        f"Generated: {ts_human()}",
        f"Scan Engine: ZAP + Ollama AI  |  Model: {CFG.model}",
        f"Scan ID: {stats.scan_run_id}",
        f"Report Classification: CONFIDENTIAL",
    ]:
        r = subtitle.add_run(line + "\n")
        r.font.size = Pt(11)
        r.font.name = "Calibri"

    doc.add_page_break()

    # ── Table of Contents (manual) ────────────────────────────────────────────
    h("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Key Scan Metrics",
        "3. Authentication Research",
        "4. Top 10 Findings by Priority",
        "5. OWASP Top 10 (2021) Mapping",
        "6. Before vs After AI Enrichment",
        "7. Scan Policy Optimization",
        "8. Trend & Pattern Analysis",
        "9. False Positive Analysis",
        "10. AI vs ZAP Severity Validation",
        "11. AI Accuracy Validation",
        "12. Remediation Recommendations",
        "Appendix A: Authentication Log",
    ]
    for item in toc_items:
        para(item, size=10)
    doc.add_page_break()

    # ── 1. Executive Summary ──────────────────────────────────────────────────
    h("1. Executive Summary")
    for paragraph in re.split(r"\n\n+", exec_summary.strip()):
        if paragraph.strip():
            para(paragraph.strip(), size=11)
    doc.add_paragraph()

    # ── 2. Key Scan Metrics ───────────────────────────────────────────────────
    h("2. Key Scan Metrics", level=2)
    high  = sum(1 for a in alerts if severity_int(
        a.get("ai_new_severity") or a.get("risk", "")) >= 3)
    med   = sum(1 for a in alerts if severity_int(
        a.get("ai_new_severity") or a.get("risk", "")) == 2)
    low   = len(alerts) - high - med

    metrics_tbl = doc.add_table(rows=9, cols=2)
    metrics_tbl.style = "Table Grid"
    metrics_pairs = [
        ("Total Unique Findings (post-deduplication)", len(alerts)),
        ("High / Critical",                            high),
        ("Medium",                                     med),
        ("Low / Informational",                        low + len(info_alerts)),
        ("Duplicates Removed by AI",                   stats.duplicates_removed),
        ("Dead Paths Pruned",                          stats.dead_paths_pruned),
        ("Auth Errors Detected (401/403)",             stats.auth_401_403_count),
        ("AI Enrichment Duration (s)",                 f"{stats.ai_duration:.1f}"),
        ("Total Scan Duration (s)",                    f"{stats.total_duration:.1f}"),
    ]
    fill_map = {1: "FFCDD2", 2: "FFE0B2", 3: "C8E6C9"}
    for i, (label, val) in enumerate(metrics_pairs):
        cell_text(metrics_tbl.rows[i].cells[0], label, bold=True)
        cell_text(metrics_tbl.rows[i].cells[1], val, align="center")
        if i in fill_map:
            shade(metrics_tbl.rows[i].cells[1], fill_map[i])
    doc.add_paragraph()

    # ── 3. Authentication Research ────────────────────────────────────────────
    h("3. Authentication Research", level=2)
    if auth_analysis:
        para(auth_analysis.get("executive_summary", ""), size=10)
        auth_tbl = add_table_with_header(["Pattern", "Freq.", "Likely Cause",
                                           "Severity", "Recommendation"])
        for fp in auth_analysis.get("failure_patterns", []):
            row_cells = auth_tbl.add_row().cells
            vals = [fp.get("pattern", ""), fp.get("frequency", ""),
                    fp.get("likely_cause", ""), fp.get("severity", ""),
                    fp.get("recommendation", "")]
            for c, v in enumerate(vals):
                cell_text(row_cells[c], v, align="center" if c in (1, 3) else "left")
            shade(row_cells[3], risk_colour_hex(fp.get("severity", "")))
    else:
        para("Authentication was disabled or no auth log was generated.", size=10)

    if auth_correlation:
        doc.add_paragraph()
        para(f"Auth overhead: ~{auth_correlation.get('estimated_auth_overhead_s', 0):.1f}s "
             f"({auth_correlation.get('auth_overhead_pct', 0):.1f}% of scan time). "
             f"Impact assessment: {auth_correlation.get('impact_assessment', 'N/A')}",
             size=10)
    doc.add_paragraph()

    # ── 4. Top 10 Findings ────────────────────────────────────────────────────
    h("4. Top 10 Findings by Priority Score", level=2)
    top10 = sorted(alerts, key=lambda a: a.get("priority_score", 0), reverse=True)[:10]
    tbl   = add_table_with_header(["#", "Name", "ZAP Risk", "AI Severity",
                                    "Priority", "OWASP", "AI Summary"])
    for idx, a in enumerate(top10, 1):
        rc = tbl.add_row().cells
        vals = [
            idx, a.get("name", ""), a.get("risk", ""),
            a.get("ai_new_severity", ""), a.get("priority_score", 0),
            f"{a.get('owasp_id','')} – {a.get('owasp_category','')}",
            (a.get("ai_summary") or "")[:180],
        ]
        for c, v in enumerate(vals):
            cell_text(rc[c], v, align="center" if c in (0, 4) else "left")
        clr = risk_colour_hex(a.get("ai_new_severity") or a.get("risk"))
        shade(rc[2], clr)
        shade(rc[3], clr)
    doc.add_paragraph()

    # ── 5. OWASP Mapping ─────────────────────────────────────────────────────
    h("5. OWASP Top 10 (2021) Mapping", level=2)
    owasp_counts = Counter(a.get("owasp_id", "A05") for a in alerts)
    tbl2 = add_table_with_header(["OWASP ID", "Category", "Count"])
    for oid, count in sorted(owasp_counts.items()):
        rc = tbl2.add_row().cells
        cell_text(rc[0], oid, bold=True, align="center")
        cell_text(rc[1], OWASP_TOP10_2021.get(oid, "Unknown"))
        cell_text(rc[2], count, align="center")
    doc.add_paragraph()

    # ── 6. Before vs After ────────────────────────────────────────────────────
    h("6. Before vs After AI Enrichment", level=2)
    ba_tbl = add_table_with_header(["Metric", "Before AI", "After AI", "Change"],
                                    header_fill="283593")
    for row_data in comparison_rows:
        rc = ba_tbl.add_row().cells
        for c, v in enumerate(row_data):
            cell_text(rc[c], v, bold=(c == 0),
                      align="center" if c > 0 else "left")
        change = str(row_data[3]) if len(row_data) > 3 else ""
        if change.startswith("-") and change not in ("-0", "-", "—"):
            shade(rc[3], "C8E6C9")
        elif change.startswith("+"):
            shade(rc[3], "FFCDD2")
    doc.add_paragraph()

    # ── 7. Scan Policy Optimization ───────────────────────────────────────────
    h("7. Scan Policy Optimization", level=2)
    if policy_report:
        para(policy_report.get("policy_summary", "No policy summary generated."), size=10)
        para(f"Estimated time reduction: {policy_report.get('estimated_time_reduction_pct', 0)}% | "
             f"Estimated noise reduction: {policy_report.get('estimated_noise_reduction_pct', 0)}%",
             bold=True, size=10)
        disable = policy_report.get("disable_test_cases", [])
        if disable:
            doc.add_paragraph()
            para("Test Cases Recommended for Disabling:", bold=True, size=10)
            policy_tbl = add_table_with_header(["Test Case", "Reason", "Confidence"],
                                               header_fill="B71C1C")
            for item in disable:
                rc = policy_tbl.add_row().cells
                cell_text(rc[0], item.get("test_case", ""))
                cell_text(rc[1], item.get("reason", ""), )
                cell_text(rc[2], item.get("confidence", ""), align="center")
    else:
        para("Policy optimization was not run or returned no proposals.", size=10)
    doc.add_paragraph()

    # ── 8. Trend Analysis ─────────────────────────────────────────────────────
    h("8. Trend & Pattern Analysis", level=2)
    if trend_report:
        para(trend_report.get("narrative", ""), size=10)
        td = trend_report.get("trend", {})
        if td and "delta_findings_vs_last" in td:
            para(f"Trend direction: {td.get('trend_direction', 'Unknown')} | "
                 f"Δ findings vs last scan: {td.get('delta_findings_vs_last', 'N/A')} | "
                 f"Avg scan duration (last 10): {td.get('avg_scan_duration_10_s', 'N/A')}s",
                 bold=True, size=10)
    else:
        para("Trend analysis was not run.", size=10)
    doc.add_paragraph()

    # ── 9. False Positive Analysis ────────────────────────────────────────────
    h("9. False Positive Analysis (Manual vs AI)", level=2)
    if fpa_report:
        para(fpa_report.get("assessment", ""), size=10)
        fpa_tbl = add_table_with_header(["Metric", "Value"], header_fill="4A148C")
        for k, v in [
            ("Manual FP Count",    fpa_report.get("manual_fp_count", 0)),
            ("AI FP Count",        fpa_report.get("ai_fp_count", 0)),
            ("Agreement",          f"{fpa_report.get('agreement_pct', 0):.1f}%"),
            ("Missed by AI",       len(fpa_report.get("missed_by_ai", []))),
            ("Over-flagged by AI", len(fpa_report.get("over_flagged_by_ai", []))),
        ]:
            rc = fpa_tbl.add_row().cells
            cell_text(rc[0], k, bold=True)
            cell_text(rc[1], v, align="center")
    else:
        para("No manual_fpa.json provided. Create this file to enable FPA comparison.", size=10)
    doc.add_paragraph()

    # ── 10. Severity Validation ───────────────────────────────────────────────
    h("10. AI vs ZAP Severity Validation", level=2)
    para(
        f"Agreement: {sev_validation.get('agreement_pct', 0):.1f}% "
        f"({sev_validation.get('agreed', 0)}/{sev_validation.get('total', 0)}). "
        f"AI escalated {sev_validation.get('escalated', 0)}, "
        f"downgraded {sev_validation.get('downgraded', 0)} findings.",
        size=10
    )
    if sev_validation.get("details"):
        dis_tbl = add_table_with_header(["Finding", "URL", "ZAP Severity", "AI Severity"],
                                         header_fill="B71C1C")
        for d in sev_validation["details"][:20]:
            rc = dis_tbl.add_row().cells
            cell_text(rc[0], d.get("name", ""))
            cell_text(rc[1], d.get("url", ""))
            cell_text(rc[2], d.get("zap", ""), align="center")
            cell_text(rc[3], d.get("ai", ""),  align="center")
            shade(rc[2], risk_colour_hex(d.get("zap")))
            shade(rc[3], risk_colour_hex(d.get("ai")))
    doc.add_paragraph()

    # ── 11. AI Validation ─────────────────────────────────────────────────────
    h("11. AI Accuracy Validation", level=2)
    if validation_report:
        para(f"Manual accuracy: {validation_report.get('manual_accuracy_pct') or 'N/A'}% | "
             f"Proxy accuracy: {validation_report.get('proxy_accuracy_pct', 0):.1f}% | "
             f"Reviewed: {validation_report.get('manually_reviewed_count', 0)} findings",
             size=10)
        para(validation_report.get("recommendations", ""), bold=True, size=10)
        para(validation_report.get("note", ""), size=9)
    doc.add_paragraph()

    # ── 12. Recommendations ──────────────────────────────────────────────────
    h("12. Remediation Recommendations", level=2)
    recs = [
        "Remediate findings in descending priority score order. Apply AI-suggested "
        "code fixes and validate each one manually before closing.",
        "Implement parameterised queries, input validation, and output encoding across "
        "all endpoints — especially those mapped to OWASP A03 (Injection).",
        "Address all High/Critical findings before the next release cycle. Use the "
        "severity validation report to review cases where AI and ZAP disagree.",
        "Apply the proposed scan policy changes to reduce noise and scan time in future runs.",
        "Batch-remediate recurring low-risk findings using shared middleware or libraries "
        "rather than fixing each URL individually.",
        "Integrate this scan into CI/CD: fail the pipeline on any finding with priority "
        "score ≥ 7 or OWASP category A01 / A03.",
        f"Populate {ValidationTracker.MANUAL_REVIEW_FILE} with analyst verdicts to improve "
        "AI accuracy tracking over time.",
        "Review manual_fpa.json comparison to tune false positive thresholds for future scans.",
    ]
    for rec in recs:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(rec)
        r.font.name = "Calibri"
        r.font.size = Pt(10)

    # ── Appendix A: Auth Log ──────────────────────────────────────────────────
    if os.path.exists(AUTH_LOG_FILE):
        doc.add_page_break()
        h("Appendix A: Authentication Log", level=2)
        para(f"Full auth log: {AUTH_LOG_FILE}", size=9)
        try:
            with open(AUTH_LOG_FILE, encoding="utf-8") as fh:
                content = fh.read()
            p = doc.add_paragraph()
            r = p.add_run(content[:4000])
            r.font.name = "Courier New"
            r.font.size = Pt(7)
        except Exception:
            pass

    doc.save(path)
    log.info("Word report → %s", path)
    return path


# ─────────────────────────────────────────────────────────────────────────────
# PROFESSIONAL HTML REPORT
# ─────────────────────────────────────────────────────────────────────────────

def generate_html_report(alerts: List[Dict], path: str = None) -> str:
    if path is None:
        path = os.path.join(CFG.report_dir, f"pro_report_{ts()}.html")

    total  = len(alerts)
    high   = sum(1 for a in alerts if severity_int(
        a.get("ai_new_severity") or a.get("risk", "")) >= 3)
    medium = sum(1 for a in alerts if severity_int(
        a.get("ai_new_severity") or a.get("risk", "")) == 2)

    SEV_CSS = {"critical": "#6a0dad", "high": "#d32f2f",
               "medium": "#f57c00", "low": "#388e3c", "unknown": "#757575"}

    def badge(sev):
        s  = (sev or "unknown").lower()
        bg = SEV_CSS.get(s, "#757575")
        return (f'<span style="background:{bg};color:#fff;padding:2px 8px;'
                f'border-radius:4px;font-size:.8em">{sev or "Unknown"}</span>')

    css = """
body{font-family:'Segoe UI',Arial,sans-serif;margin:0;color:#212121;background:#fafafa}
.header{background:linear-gradient(135deg,#1a237e,#3949ab);color:#fff;padding:32px 40px}
.header h1{margin:0;font-size:2em;font-weight:700}.header p{margin:4px 0;opacity:.85}
.container{max-width:1400px;margin:0 auto;padding:32px 40px}
.kpi-grid{display:grid;grid-template-columns:repeat(4,1fr);gap:20px;margin:24px 0}
.kpi{background:#fff;border-radius:12px;padding:20px;text-align:center;
     box-shadow:0 2px 8px rgba(0,0,0,.08)}
.kpi .num{font-size:2.2em;font-weight:700;line-height:1}
.kpi .lbl{font-size:.85em;color:#666;margin-top:6px}
h2{color:#1a237e;border-bottom:2px solid #e8eaf6;padding-bottom:6px;margin-top:36px}
table{border-collapse:collapse;width:100%;font-size:.88em;background:#fff;
      border-radius:8px;overflow:hidden;box-shadow:0 2px 8px rgba(0,0,0,.05)}
th{background:#3949ab;color:#fff;padding:10px 8px;text-align:left;font-weight:600}
td{border-bottom:1px solid #f0f0f0;padding:8px;vertical-align:top}
tr:hover{background:#f3f4fc}.score{font-weight:700;font-size:1.1em}
pre{background:#f5f5f5;padding:8px;border-radius:4px;overflow-x:auto;font-size:.82em;margin:4px 0}
.badge-recurring{background:#7b1fa2;color:#fff;padding:1px 6px;border-radius:4px;font-size:.8em}
footer{text-align:center;padding:24px;color:#999;font-size:.8em;border-top:1px solid #e0e0e0;margin-top:40px}
"""
    sorted_alerts = sorted(alerts, key=lambda a: a.get("priority_score", 0), reverse=True)

    rows_html = ""
    for i, a in enumerate(sorted_alerts, 1):
        rem   = a.get("ai_remediation_snippet", {}) or {}
        score = a.get("priority_score", 0)
        score_clr = "#d32f2f" if score >= 7 else "#f57c00" if score >= 4 else "#388e3c"
        rec_cell = (f"<b>{rem.get('language','')}</b>: {rem.get('explanation','')}"
                    f"<br><pre>{str(rem.get('code',''))[:400]}</pre>"
                    if rem.get("code") else "—")
        recurring = ('<span class="badge-recurring">↻ Recurring</span>'
                     if a.get("is_recurring") else f"×{a.get('occurrence_count', 1)}")
        rows_html += f"""<tr>
<td>{i}</td>
<td>{a.get('name','')}</td>
<td style="word-break:break-all;font-size:.8em">{a.get('url','')}</td>
<td>{a.get('param') or ''}</td>
<td>{badge(a.get('risk'))}</td>
<td>{badge(a.get('ai_new_severity'))}</td>
<td><span class="score" style="color:{score_clr}">{score}</span></td>
<td>{a.get('owasp_id','')} {a.get('owasp_category','')}</td>
<td>{a.get('exploitation_rank','')}</td>
<td>{(a.get('ai_summary') or '')[:250]}</td>
<td>{rec_cell}</td>
<td>{recurring}</td>
<td>{'✔' if a.get('severity_match') else '✘'} {a.get('severity_delta','')}</td>
</tr>"""

    html = f"""<!doctype html><html lang="en"><head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>DAST Security Report — {CFG.target}</title>
<style>{css}</style></head><body>
<div class="header">
  <h1>🛡️ DAST Security Assessment Report</h1>
  <p><strong>Target:</strong> {CFG.target}</p>
  <p><strong>Generated:</strong> {ts_human()} &nbsp;|&nbsp;
     <strong>Scan ID:</strong> {STATS.scan_run_id} &nbsp;|&nbsp;
     <strong>Engine:</strong> ZAP + Ollama ({CFG.model})</p>
</div>
<div class="container">
<div class="kpi-grid">
  <div class="kpi"><div class="num" style="color:#1a237e">{total}</div>
    <div class="lbl">Total Findings</div></div>
  <div class="kpi"><div class="num" style="color:#d32f2f">{high}</div>
    <div class="lbl">High / Critical</div></div>
  <div class="kpi"><div class="num" style="color:#f57c00">{medium}</div>
    <div class="lbl">Medium</div></div>
  <div class="kpi"><div class="num" style="color:#388e3c">{total-high-medium}</div>
    <div class="lbl">Low / Other</div></div>
</div>
<h2>Findings — Sorted by Priority Score ▼</h2>
<table><thead><tr>
  <th>#</th><th>Name</th><th>URL</th><th>Param</th>
  <th>ZAP Risk</th><th>AI Severity</th><th>Score ▼</th>
  <th>OWASP</th><th>Exploit Rank</th><th>AI Summary</th>
  <th>Remediation</th><th>Recurring</th><th>Sev Delta</th>
</tr></thead><tbody>
{rows_html}
</tbody></table>
<h2>Recommendations</h2>
<ol>
  <li>Remediate by priority score (highest first) — apply AI code fixes and validate manually.</li>
  <li>Address all OWASP A03 Injection findings immediately.</li>
  <li>Apply proposed scan policy changes before the next scan run.</li>
  <li>Fail CI/CD pipeline on any finding with priority score ≥ 7.</li>
</ol>
</div>
<footer>
  Generated by ZAP AI DAST Agent v4 &nbsp;|&nbsp; {ts_human()} &nbsp;|&nbsp;
  Scan ID: {STATS.scan_run_id} &nbsp;|&nbsp; CONFIDENTIAL
</footer>
</body></html>"""

    with open(path, "w", encoding="utf-8") as f:
        f.write(html)
    log.info("HTML report → %s", path)
    return path


def generate_xml_report(zap, alerts: List[Dict], path: str = None) -> str:
    if path is None:
        path = os.path.join(CFG.report_dir, f"report_{ts()}.xml")
    if zap:
        try:
            xml_content = zap.core.xmlreport()
            if xml_content:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(xml_content)
                log.info("ZAP XML report → %s", path)
                return path
        except Exception:
            pass

    suite = ET.Element("testsuite", name="dast", tests=str(len(alerts)))
    for a in alerts:
        tc  = ET.SubElement(suite, "testcase",
                            classname=a.get("name", "alert"),
                            name=a.get("url", "unknown"))
        sev = a.get("ai_new_severity") or a.get("risk") or "Unknown"
        if severity_int(sev) >= 3 or a.get("priority_score", 0) >= 7:
            f = ET.SubElement(tc, "failure",
                              message=f"Severity:{sev} Score:{a.get('priority_score',0)}")
            f.text = (a.get("ai_summary") or "")[:1000]
        else:
            ET.SubElement(tc, "system-out").text = (a.get("ai_summary") or "")[:1000]

    ET.ElementTree(suite).write(path, encoding="utf-8", xml_declaration=True)
    log.info("XML report → %s", path)
    return path


# ─────────────────────────────────────────────────────────────────────────────
# MAIN PIPELINE
# ─────────────────────────────────────────────────────────────────────────────

def run() -> None:
    log.info("╔════════════════════════════════════════╗")
    log.info("║  ZAP AI DAST Agent v4  — Starting      ║")
    log.info("╚════════════════════════════════════════╝")
    log.info("Target: %s | Scan ID: %s", CFG.target, STATS.scan_run_id)

    STATS.scan_start = time.time()
    load_ai_cache()

    # ── 1. Start ZAP ─────────────────────────────────────────────────────────
    try:
        start_zap()
        wait_for_zap_api(timeout=180)
    except Exception as e:
        log.error("ZAP startup failed: %s", e)
        log_auth_failure("ZAP Startup", "ZAP failed to start", e)
        return

    zap = zap_client()

    # ── 2. Auth setup ─────────────────────────────────────────────────────────
    context_name = context_id = user_id = None
    auth_ok      = False
    username = password = None

    if CFG.enable_auth:
        auth_result  = run_full_auth_setup()
        auth_ok      = auth_result["success"]
        context_id   = auth_result["context_id"]
        user_id      = auth_result["user_id"]
        username     = auth_result["username"]
        password     = auth_result["password"]
        context_name = "vulnweb" if context_id else None
        if not auth_ok:
            log.warning("Auth incomplete — scans may be unauthenticated")

    # ── PILLAR 1: Detect 401/403 sequences (run after auth is set up) ─────────
    seq_analysis: Dict = {}
    if CFG.enable_auth or True:  # always check if ZAP is active
        seq_analysis = detect_401_403_sequences(zap)

    # ── 3. Discover & Spider ──────────────────────────────────────────────────
    zap_force_discover(zap, CFG.target)
    raw_urls = spider_target(zap, CFG.target,
                             context_name=context_name if auth_ok else None,
                             user_id=user_id if auth_ok else None)

    # ── 4. Scope filtering ────────────────────────────────────────────────────
    filtered_urls = filter_urls(raw_urls)

    # ── PILLAR 4: AI dead path reduction ──────────────────────────────────────
    log.info("[Pillar 4] AI dead path / crawl scope reduction…")
    urls, dead_urls = reduce_crawl_scope(filtered_urls)
    log.info("Final scan scope: %d URLs", len(urls))

    # ── 5. Active scans with session watchdog ─────────────────────────────────
    stop_event = threading.Event()
    watchdog   = None
    if auth_ok and username and password:
        watchdog = SessionWatchdog(username, password, stop_event)
        watchdog.start()

    if CFG.enable_active_scan:
        run_parallel_scans(zap, urls,
                           context_name=context_name if auth_ok else None,
                           user_id=user_id if auth_ok else None)

    if watchdog:
        stop_event.set()
        watchdog.join(timeout=10)

    # ── 6. Collect alerts ─────────────────────────────────────────────────────
    log.info("Collecting ZAP alerts…")
    try:
        raw_alerts = zap.core.alerts(baseurl=CFG.target) if zap else []
    except Exception as e:
        log.error("Alert collection failed: %s", e)
        raw_alerts = []

    seen, unique = set(), []
    for a in raw_alerts:
        key = (a.get("alertRef"), a.get("url"), a.get("param"))
        if key not in seen and a.get("risk", "").lower() in CFG.allowed_risks:
            seen.add(key)
            unique.append(a)
    log.info("Unique in-risk alerts: %d", len(unique))

    # ── Record BEFORE stats ───────────────────────────────────────────────────
    STATS.before_ai_count  = len(unique)
    STATS.before_ai_high   = sum(1 for a in unique if severity_int(a.get("risk", "")) >= 3)
    STATS.before_ai_medium = sum(1 for a in unique if severity_int(a.get("risk", "")) == 2)
    write_csv(unique, os.path.join(CFG.report_dir, f"before_ai_{ts()}.csv"))

    # ── PILLAR 2: False positive reduction ────────────────────────────────────
    log.info("[Pillar 2] AI duplicate grouping…")
    unique = ai_group_duplicates(unique)

    log.info("[Pillar 2] Extracting informational findings…")
    actionable, info_alerts = extract_information_findings(unique)

    log.info("[Pillar 2] Identifying recurring findings…")
    actionable = identify_recurring_findings(actionable)

    # ── PILLAR 3: AI enrichment + ranking + OWASP ─────────────────────────────
    STATS.ai_enrichment_start = time.time()
    enriched = actionable
    if CFG.enable_ai_analysis:
        try:
            enriched = analyze_alerts(actionable)
        except Exception as e:
            log.error("AI enrichment failed: %s", e)

    log.info("[Pillar 3] AI exploitation & impact ranking…")
    enriched = ai_rank_findings(enriched)

    log.info("[Pillar 3] OWASP Top 10 mapping…")
    enriched = map_owasp_top10(enriched)
    STATS.ai_enrichment_end = time.time()
    save_ai_cache()

    # ── Priority scores ───────────────────────────────────────────────────────
    for a in enriched:
        a["priority_score"] = compute_priority_score(a)

    # ── PILLAR 3: Severity validation ─────────────────────────────────────────
    log.info("[Pillar 3] AI vs ZAP severity validation…")
    sev_validation = validate_ai_vs_zap_severity(enriched)

    # ── Record AFTER stats ────────────────────────────────────────────────────
    STATS.after_ai_count  = len(enriched)
    STATS.after_ai_high   = sum(1 for a in enriched
                                if severity_int(a.get("ai_new_severity") or a.get("risk", "")) >= 3)
    STATS.after_ai_medium = sum(1 for a in enriched
                                if severity_int(a.get("ai_new_severity") or a.get("risk", "")) == 2)
    STATS.scan_end = time.time()

    write_csv(enriched, os.path.join(CFG.report_dir, f"after_ai_{ts()}.csv"))

    # ── PILLAR 1: Auth log AI analysis ────────────────────────────────────────
    log.info("[Pillar 1] AI analysis of authentication log…")
    auth_analysis   = analyze_auth_log_with_ai()
    auth_correlation = correlate_auth_with_duration(auth_analysis, seq_analysis)

    # ── PILLAR 2: Manual vs AI FPA comparison ────────────────────────────────
    log.info("[Pillar 2] Manual vs AI false positive comparison…")
    fpa_report = compare_manual_vs_ai_fpa(enriched + info_alerts)

    # ── PILLAR 4: Scan policy optimization ────────────────────────────────────
    log.info("[Pillar 4] Scan policy optimization…")
    policy_optimizer = ScanPolicyOptimizer()
    policy_report    = policy_optimizer.propose_policy_changes(enriched, dead_urls, urls)

    # ── PILLAR 5: Trend analysis ──────────────────────────────────────────────
    log.info("[Pillar 5] Trend & pattern analysis…")
    trend_analyzer = TrendAnalyzer()
    trend_report   = trend_analyzer.analyze(STATS, enriched)

    # ── PILLAR 7: AI validation ────────────────────────────────────────────────
    log.info("[Pillar 7] AI accuracy validation…")
    validation_report = ValidationTracker().evaluate(enriched, sev_validation)

    # ── PILLAR 6: Comparison tables + reports ────────────────────────────────
    comparison_rows = build_comparison_rows(STATS, sev_validation)
    write_comparison_csv(
        comparison_rows,
        os.path.join(CFG.report_dir, f"comparison_{ts()}.csv")
    )

    # Summary log
    log.info(
        "═══ SCAN METRICS ═══\n"
        "  Before AI : %d findings (%d high, %d medium)\n"
        "  After AI  : %d findings (%d high, %d medium)\n"
        "  Duplicates removed : %d | Dead paths pruned: %d\n"
        "  Auth errors (401/403): %d\n"
        "  AI enrichment: %.1fs | Total scan: %.1fs",
        STATS.before_ai_count, STATS.before_ai_high, STATS.before_ai_medium,
        STATS.after_ai_count,  STATS.after_ai_high,  STATS.after_ai_medium,
        STATS.duplicates_removed, STATS.dead_paths_pruned,
        STATS.auth_401_403_count,
        STATS.ai_duration, STATS.total_duration,
    )

    top5 = sorted(enriched, key=lambda x: x.get("priority_score", 0), reverse=True)[:5]
    log.info("Top-5 findings:")
    for a in top5:
        log.info("  [%.1f] %s — %s", a["priority_score"], a.get("name"), a.get("url"))

    # ── Reports ───────────────────────────────────────────────────────────────
    if CFG.enable_report:
        try:
            generate_html_report(enriched)
        except Exception as e:
            log.error("HTML report failed: %s", e)

        try:
            generate_xml_report(zap, enriched)
        except Exception as e:
            log.error("XML report failed: %s", e)

        try:
            if zap:
                zap_html = zap.core.htmlreport()
                p = os.path.join(CFG.report_dir, f"zap_native_{ts()}.html")
                with open(p, "w", encoding="utf-8") as f:
                    f.write(zap_html)
                log.info("ZAP native HTML → %s", p)
        except Exception as e:
            log.warning("ZAP native HTML failed: %s", e)

        try:
            generate_excel_report(
                enriched, info_alerts, sev_validation, comparison_rows,
                fpa_report, policy_report, trend_report, validation_report,
            )
        except Exception as e:
            log.error("Excel report failed: %s", e)

        try:
            generate_word_report(
                enriched, info_alerts, sev_validation, comparison_rows,
                STATS, auth_analysis, auth_correlation,
                fpa_report, policy_report, trend_report, validation_report,
            )
        except Exception as e:
            log.error("Word report failed: %s", e)

    save_ai_cache()

    log.info("╔════════════════════════════════════════╗")
    log.info("║  SCAN COMPLETE                          ║")
    log.info("║  Reports → %s/", CFG.report_dir)
    log.info("╚════════════════════════════════════════╝")


if __name__ == "__main__":
    run()
